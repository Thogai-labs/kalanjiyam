from collections.abc import Iterator
from datetime import date
import io
import json
import os
import re
import zipfile

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches

from kalanjiyam.utils.page_document import PageDocument, document_for_revision
from kalanjiyam.utils.storage import get_storage, editor_image_key

DOUBLE_DANDA = "\u0965"

TEI_HEADER_BOILERPLATE = """
<?xml version="1.0" encoding="UTF-8"?>
<!-- This file was automatically generated. Please review it for markup mistakes
and resolve any TODOs. -->
<TEI xml:id="{xml_id}" xmlns="http://www.tei-c.org/ns/1.0">
  <teiHeader xml:lang="en">
    <fileDesc>
      <titleStmt>
        <title type="main">{title}</title>
        <title type="sub">A machine-readable edition</title>
        <author>{author}</author>
      </titleStmt>
      <publicationStmt>
        <publisher>Kalanjiyam</publisher>
        <!-- "free" or "restricted" depending on the license-->
        <availability status="{availability_status}">
          <license>
            TODO
          </license>
        </availability>
        <date>{current_year}</date>
      </publicationStmt>
      <sourceDesc>
        <bibl>
          <title>{title}</title>
          <author>{author}</author>
          <editor>{editor}</editor>
          <publisher>{publisher}</publisher>
          <pubPlace>{publisher_location}</pubPlace>
          <date>{publication_year}</date>
        </bibl>
      </sourceDesc>
    </fileDesc>
    <encodingDesc>
      <projectDesc>
        <p>Produced through the distributed proofreading interface on Kalanjiyam.</p>
      </projectDesc>
    </encodingDesc>
    <revisionDesc>
      TODO
    </revisionDesc>
  </teiHeader>
  <text xml:lang="{text_language}">
    <body>
""".strip()

PageContent = str
Line = str


def _iter_raw_text_lines(blobs: list[PageContent]) -> Iterator[Line]:
    """Iterate over text blobs as a stream of lines."""
    for blob in blobs:
        blob = blob.strip()
        for line in blob.splitlines():
            yield line.strip()


def iter_blocks(blobs: Iterator[PageContent]) -> Iterator[list[Line]]:
    """Iterate over text blobs as a stream of blocks.

    A block is a sequence of lines separated by an empty line."""
    buf = []
    for line in _iter_raw_text_lines(blobs):
        if line:
            buf.append(line)
        elif buf:
            yield buf
            buf = []
    if buf:
        yield buf


def is_verse(lines: list[Line]) -> bool:
    """Heuristically decide whether a list of lines represents a verse."""
    return lines[-1].endswith(DOUBLE_DANDA)


def create_plain_text_block(lines: list[Line]) -> str:
    """Convert a group of lines into a well-formatted plain-text block."""
    if is_verse(lines):
        return "\n".join(lines)

    buf = []
    for line in lines:
        # Join hyphens
        if line.endswith("-"):
            buf.append(line[:-1])
        else:
            buf.append(line)
            buf.append(" ")
    return "".join(buf).strip()


def create_tei_header_boilerplate(**kw) -> str:
    # FIXME: add much more TEI boilerplate
    return TEI_HEADER_BOILERPLATE.format(**kw)


def create_xml_block(lines: list[Line]) -> str:
    """Convert a group of lines into a well-formatted TEI XML block."""
    if is_verse(lines):
        buf = ["<lg>"]
        for line in lines:
            buf.append(f"  <l>{line}</l>")
        buf.append("</lg>")
        return "\n".join(buf)

    buf = ["<p>"]
    for line in lines:
        # Join hyphens
        if line.endswith("-"):
            buf.append(line[:-1])
        else:
            buf.append(line)
            buf.append(" ")

    # Strip trailing space from the loop.
    buf[-1] = buf[-1].strip()

    buf.append("</p>")
    return "".join(buf).strip()


def to_plain_text(blobs: list[PageContent]) -> str:
    """Publish a project as plain text."""
    blocks = iter_blocks(blobs)
    return "\n\n".join(create_plain_text_block(b) for b in blocks)


def to_tei_xml(project_meta: dict[str, str], blobs: list[PageContent]) -> str:
    """Publish a project as TEI XML."""
    project_meta.update(
        {
            "xml_id": "TODO",
            "current_year": date.today().year,
            "publisher_location": "TODO",
            "text_language": "sa-Deva",
            # "free" or "restricted"
            "availability_status": "TODO",
        }
    )
    buf = [create_tei_header_boilerplate(**project_meta)]

    for i, blob in enumerate(blobs):
        page_number = i + 1
        buf.append(f'<pb n="{page_number}" />')

        # <pb> element makes it difficult to work with a stream of blobs,
        # so just process one blob at a time and stitch them together after.
        blocks = iter_blocks([blob])
        buf.append("\n\n".join(create_xml_block(b) for b in blocks))

    buf.append("</body></text></TEI>")
    return "\n\n".join(buf)


def revision_plain_content(revision) -> str:
    """Plain text for a revision, preferring structured document when present."""
    if revision is None:
        return ""
    if getattr(revision, "document", None):
        return PageDocument.from_dict(revision.document).to_plain_text()
    return revision.content or ""


def document_to_tei(doc: PageDocument) -> str:
    return doc.to_tei_fragment()


def documents_to_tei_xml(project_meta: dict[str, str], pages) -> str:
    """TEI XML from pages with structured revisions when available."""
    project_meta.update(
        {
            "xml_id": "TODO",
            "current_year": date.today().year,
            "publisher_location": "TODO",
            "text_language": "sa-Deva",
            "availability_status": "TODO",
        }
    )
    buf = [create_tei_header_boilerplate(**project_meta)]
    for i, page in enumerate(pages):
        page_number = i + 1
        buf.append(f'<pb n="{page_number}" />')
        if page.revisions:
            rev = page.revisions[-1]
            if getattr(rev, "document", None) and rev.content_format == "blocks":
                doc = PageDocument.from_dict(rev.document)
                buf.append(doc.to_tei_fragment())
            else:
                blocks = iter_blocks([rev.content or ""])
                buf.append("\n\n".join(create_xml_block(b) for b in blocks))
    buf.append("</body></text></TEI>")
    return "\n\n".join(buf)


def documents_to_html(pages, *, replica: bool = False) -> str:
    parts = [
        '<!DOCTYPE html><html><head><meta charset="utf-8">',
        "<title>Export</title>",
        '<link rel="stylesheet" href="/static/css/style.css">',
        '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.27/dist/katex.min.css">',
        '<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.27/dist/katex.min.js"></script>',
        '<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.27/dist/contrib/auto-render.min.js" onload="renderMathInElement(document.body, {delimiters:[{left:\'$$\',right:\'$$\',display:true},{left:\'$\',right:\'$\',display:false},{left:\'\\\\(\',right:\'\\\\)\',display:false},{left:\'\\\\[\',right:\'\\\\]\',display:true}],throwOnError:false});"></script>',
        "</head><body class='p-8'>",
    ]
    for page in pages:
        parts.append(f'<section class="ocr-export-page" data-page="{page.slug}">')
        if page.revisions:
            rev = page.revisions[-1]
            if getattr(rev, "document", None):
                doc = document_for_revision(rev, page)
                parts.append(doc.to_html(replica=replica))
            else:
                parts.append(f"<pre>{rev.content}</pre>")
        parts.append("</section>")
    parts.append("</body></html>")
    return "\n".join(parts)


def documents_to_json_bundle(project, pages) -> str:
    payload = {
        "format_version": "3.0",
        "project_slug": project.slug,
        "display_title": project.display_title,
        "pages": [],
    }
    for page in pages:
        entry = {
            "slug": page.slug,
            "order": page.order,
            "page_width": page.page_width,
            "page_height": page.page_height,
            "ocr_bounding_boxes": page.ocr_bounding_boxes,
        }
        if page.revisions:
            rev = page.revisions[-1]
            entry["revision"] = {
                "content": rev.content,
                "content_format": getattr(rev, "content_format", "plain"),
                "document": getattr(rev, "document", None),
            }
        payload["pages"].append(entry)
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _parse_dimension(val_str: str):
    if not val_str:
        return None
    val_str = val_str.strip().lower()
    import re
    m = re.match(r"^([\d.]+)\s*(px|in|cm|pt|%)?$", val_str)
    if not m:
        return None
    num = float(m.group(1))
    unit = m.group(2) or "px"
    
    from docx.shared import Inches, Pt, Cm
    if unit == "px":
        return Pt(num * 0.75)
    elif unit == "in":
        return Inches(num)
    elif unit == "cm":
        return Cm(num)
    elif unit == "pt":
        return Pt(num)
    return None


def _add_image_to_paragraph(img_tag, paragraph) -> bool:
    src = img_tag.get("src")
    if not src:
        return False

    # Extract project_slug and filename
    match = re.search(r"static/uploads/([^/]+)/images/([^/?]+)", src)
    if not match:
        match = re.search(r"static/uploads/([^/]+)/([^/?]+)", src)

    if match:
        project_slug = match.group(1)
        filename = match.group(2)
        key = editor_image_key(project_slug, filename)

        storage = get_storage()
        try:
            if storage.exists(key):
                img_bytes = storage.read_bytes(key)
                img_io = io.BytesIO(img_bytes)

                # Parse dimensions from attributes or inline styles
                width_str = img_tag.get("width")
                height_str = img_tag.get("height")

                style = img_tag.get("style", "")
                if style:
                    w_match = re.search(r"width\s*:\s*([^;]+)", style, re.IGNORECASE)
                    if w_match:
                        width_str = w_match.group(1)
                    h_match = re.search(r"height\s*:\s*([^;]+)", style, re.IGNORECASE)
                    if h_match:
                        height_str = h_match.group(1)

                w_docx = _parse_dimension(width_str)
                h_docx = _parse_dimension(height_str)

                run = paragraph.add_run()
                if w_docx and h_docx:
                    run.add_picture(img_io, width=w_docx, height=h_docx)
                elif w_docx:
                    run.add_picture(img_io, width=w_docx)
                elif h_docx:
                    run.add_picture(img_io, height=h_docx)
                else:
                    run.add_picture(img_io, width=Inches(4.5))
                return True
        except Exception:
            # Fallback text if image load/parse failed
            paragraph.add_run(f" [Image: {filename} failed to load] ")
            return False

    alt = img_tag.get("alt", "Image")
    paragraph.add_run(f" [{alt}: {src}] ")
    return False


def auto_wrap_math(text: str) -> str:
    if not text:
        return ""
    
    import re

    # Step 1: Temporarily extract HTML tags (<img ...>, <a ...>, etc.) to protect them
    html_tags = []
    def _extract_tag(m):
        idx = len(html_tags)
        html_tags.append(m.group(0))
        return f"__TEMP_HTML_TAG_{idx}__"

    processed_text = re.sub(r'<[^>]+>', _extract_tag, text)
    parts = re.split(r'(\$\$.*?\$\$|\$.*?\$)', processed_text)
    
    def wrap_segment(seg: str) -> str:
        if not seg.strip():
            return seg
            
        trimmed = seg.strip()
        if (
            re.search(r'\.(png|jpe?g|gif|webp|svg)(\?.*)?$', trimmed, re.I)
            or re.match(r'^https?://', trimmed, re.I)
            or trimmed.startswith('/static/')
            or 'extracted_' in trimmed
            or '__TEMP_HTML_TAG_' in trimmed
        ):
            return seg
            
        has_devanagari = bool(re.search(r'[\u0900-\u097F]', seg))
        
        if not has_devanagari:
            has_latex = bool(re.search(r'\\(begin|end|Delta|times|therefore|vmatrix|frac|alpha|beta|gamma|theta|approx|neq|pm|lambda|sigma|pi|phi|omega|sqrt|partial|nabla|int|sum|prod|cup|cap|in|subset|infty|left|right|vmatrix|matrix|align|circ|text|deg)', seg))
            words = re.findall(r'[a-z]{4,}', seg)
            is_standalone_equation = len(words) == 0 or '\\begin' in seg or ' ' not in seg
            
            if is_standalone_equation:
                has_math_sub_super = bool(re.search(r'[a-zA-Z0-9]*_[a-zA-Z0-9\{\}\\\s]+|[a-zA-Z0-9]*\^[a-zA-Z0-9\{\}\\\s]+', seg))
                has_math_equation = bool(re.search(r'[a-zA-Z0-9]+\s*[\+\*=]\s*[a-zA-Z0-9]+', seg) or re.search(r'[a-zA-Z0-9]+\s+[\-\/]\s+[a-zA-Z0-9]+', seg))
                if has_latex or has_math_sub_super or has_math_equation:
                    if (
                        'extracted_' in seg
                        or '.png' in seg
                        or '.jpg' in seg
                        or '.jpeg' in seg
                        or '.svg' in seg
                        or bool(re.search(r'/[a-zA-Z0-9_-]+/', seg))
                        or bool(re.search(r'\b[a-zA-Z]:\\|\\Users\\', seg))
                        or 'http:' in seg
                        or 'https:' in seg
                        or 'src=' in seg
                        or '__TEMP_HTML_TAG_' in seg
                    ):
                        return seg
                    if '\\begin' in seg or '\n' in seg:
                        return f"$${seg}$$"
                    return f"${seg}$"
                return seg
            else:
                def token_replacer(match):
                    m = match.group(1)
                    t = m.strip()
                    if not t:
                        return m
                    if bool(re.search(r'/[a-zA-Z0-9_-]+/', t)) or '\\Users' in t or ':\\' in t or 'extracted_' in t or '__TEMP_HTML_TAG_' in t:
                        return m
                    leading_space = re.match(r'^\s*', m).group(0)
                    trailing_space = re.search(r'\s*$', m).group(0)
                    return f"{leading_space}${t}${trailing_space}"
                
                pattern = r'([a-zA-Z0-9]*[\^\_\\][a-zA-Z0-9\{\}\\\:\.\,\-\+\*\/]*[a-zA-Z0-9\}]+)'
                return re.sub(pattern, token_replacer, seg)
        else:
            result = seg
            result = re.sub(r'(\\begin\{[a-zA-Z]+\}[\s\S]*?\\end\{[a-zA-Z]+\})', r'$$\1$$', result)
            
            def devanagari_replacer(match):
                m = match.group(1)
                t = m.strip()
                if not t:
                    return m
                if bool(re.search(r'/[a-zA-Z0-9_-]+/', t)) or bool(re.search(r'\b[a-zA-Z]:\\|\\Users\\', t)) or 'extracted_' in t or '__TEMP_HTML_TAG_' in t or '.png' in t or '.jpg' in t:
                    return m
                if re.match(r'^[a-zA-Z]$', t) or re.match(r'^\d+$', t):
                    return m
                leading_space = re.match(r'^\s*', m).group(0)
                trailing_space = re.search(r'\s*$', m).group(0)
                return f"{leading_space}${t}${trailing_space}"
                
            math_pattern = r'((?:[a-zA-Z0-9\(\)\[\]\s=\+\-\*\/]*?)(?:(?:\\[a-zA-Z]+)|(?:[a-zA-Z0-9]+(?:_[a-zA-Z0-9]+|\^[a-zA-Z0-9]+)))(?:[a-zA-Z0-9\+\-\*\/=\(\)\[\]_\^\\\{\}\:\.,\s\times\therefore]|(?:\\[a-zA-Z]+)|(?:[a-zA-Z0-9]+(?:_[a-zA-Z0-9]+|\^[a-zA-Z0-9]+)))*)'
            return re.sub(math_pattern, devanagari_replacer, result)
            
    processed_parts = []
    for part in parts:
        if not part:
            continue
        if (part.startswith('$$') and part.endswith('$$')) or (part.startswith('$') and part.endswith('$')):
            processed_parts.append(part)
        else:
            processed_parts.append(wrap_segment(part))
            
    result_str = "".join(processed_parts)
    for idx, tag in enumerate(html_tags):
        result_str = result_str.replace(f"__TEMP_HTML_TAG_{idx}__", tag)

    return result_str


def _add_text_and_math_to_paragraph(text, paragraph, bold=False, italic=False, underline=False, strike=False, font_name=None, font_size=None) -> None:
    import re
    from docx.oxml import parse_xml
    import latex2mathml.converter
    import mathml2omml

    parts = re.split(r'(\$\$.*?\$\$|\$.*?\$)', text)
    for part in parts:
        if not part:
            continue
        if (part.startswith('$$') and part.endswith('$$')) or (part.startswith('$') and part.endswith('$')):
            is_display = part.startswith('$$')
            latex_formula = part[2:-2] if is_display else part[1:-1]
            latex_formula = latex_formula.strip()
            if not latex_formula:
                continue
            try:
                mathml = latex2mathml.converter.convert(latex_formula)
                omml = mathml2omml.convert(mathml)
                xml_str = f'<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml}</m:oMathPara>'
                omathpara_el = parse_xml(xml_str.encode("utf-8"))
                paragraph._element.append(omathpara_el)
            except Exception:
                run = paragraph.add_run(part)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                if strike:
                    run.font.strike = True
                if font_name:
                    run.font.name = font_name
                    try:
                        from docx.oxml.ns import qn
                        rPr = run._r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:ascii'), font_name)
                        rFonts.set(qn('w:hAnsi'), font_name)
                        rFonts.set(qn('w:eastAsia'), font_name)
                        rFonts.set(qn('w:cs'), font_name)
                    except Exception:
                        pass
                if font_size:
                    from docx.shared import Pt
                    run.font.size = Pt(font_size)
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if strike:
                run.font.strike = True
            if font_name:
                run.font.name = font_name
                try:
                    from docx.oxml.ns import qn
                    rPr = run._r.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:ascii'), font_name)
                    rFonts.set(qn('w:hAnsi'), font_name)
                    rFonts.set(qn('w:eastAsia'), font_name)
                    rFonts.set(qn('w:cs'), font_name)
                except Exception:
                    pass
            if font_size:
                from docx.shared import Pt
                run.font.size = Pt(font_size)


def _get_or_create_footnotes_part(doc):
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    
    footnotes_uri = PackURI('/word/footnotes.xml')
    footnotes_content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'
    
    package = doc.part.package
    for part in package.parts:
        if part.partname == footnotes_uri:
            return part
            
    # If not found, create empty footnotes xml boilerplate
    empty_footnotes_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
</w:footnotes>
""".strip()

    footnotes_part = Part(
        footnotes_uri,
        footnotes_content_type,
        empty_footnotes_xml.encode('utf-8'),
        package
    )
    package.parts.append(footnotes_part)
    doc.part.relate_to(
        footnotes_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    )
    return footnotes_part


def _parse_inline_elements(parent_el, paragraph, docx_doc, bold=False, italic=False, underline=False, strike=False, font_name=None, font_size=None) -> None:
    for child in parent_el.children:
        if isinstance(child, NavigableString):
            text = str(child).replace("\r", "").replace("\n", " ")
            if text:
                wrapped_text = auto_wrap_math(text)
                _add_text_and_math_to_paragraph(wrapped_text, paragraph, bold, italic, underline, strike, font_name, font_size)
        else:
            tag = child.name
            if tag == "br":
                paragraph.add_run().add_break()
            elif tag in ["strong", "b"]:
                _parse_inline_elements(child, paragraph, docx_doc, bold=True, italic=italic, underline=underline, strike=strike, font_name=font_name, font_size=font_size)
            elif tag in ["em", "i"]:
                _parse_inline_elements(child, paragraph, docx_doc, bold=bold, italic=True, underline=underline, strike=strike, font_name=font_name, font_size=font_size)
            elif tag in ["u"]:
                _parse_inline_elements(child, paragraph, docx_doc, bold=bold, italic=italic, underline=True, strike=strike, font_name=font_name, font_size=font_size)
            elif tag in ["s", "strike", "del"]:
                _parse_inline_elements(child, paragraph, docx_doc, bold=bold, italic=italic, underline=underline, strike=True, font_name=font_name, font_size=font_size)
            elif tag == "a":
                href = child.get("href", "")
                text_content = child.get_text() or href
                if href:
                    import docx
                    from docx.oxml.shared import OxmlElement, qn
                    from docx.shared import RGBColor
                    
                    part = paragraph.part
                    r_id = part.relate_to(href, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)
                    
                    new_run = docx.text.run.Run(OxmlElement('w:r'), paragraph)
                    new_run.text = text_content
                    new_run.bold = bold
                    new_run.italic = italic
                    new_run.underline = True
                    try:
                        new_run.font.color.rgb = RGBColor(5, 99, 193)
                    except Exception:
                        pass
                    if font_name:
                        new_run.font.name = font_name
                        try:
                            rPr = new_run._r.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:ascii'), font_name)
                            rFonts.set(qn('w:hAnsi'), font_name)
                            rFonts.set(qn('w:eastAsia'), font_name)
                            rFonts.set(qn('w:cs'), font_name)
                        except Exception:
                            pass
                    if font_size:
                        from docx.shared import Pt
                        new_run.font.size = Pt(font_size)
                    hyperlink.append(new_run._element)
                    paragraph._p.append(hyperlink)
                else:
                    run = paragraph.add_run(text_content)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    if strike:
                        run.font.strike = True
                    if font_name:
                        run.font.name = font_name
                        try:
                            from docx.oxml.ns import qn
                            rPr = run._r.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:ascii'), font_name)
                            rFonts.set(qn('w:hAnsi'), font_name)
                            rFonts.set(qn('w:eastAsia'), font_name)
                            rFonts.set(qn('w:cs'), font_name)
                        except Exception:
                            pass
                    if font_size:
                        from docx.shared import Pt
                        run.font.size = Pt(font_size)
            elif tag == "img":
                _add_image_to_paragraph(child, paragraph)
            elif tag == "span" and ("math-placeholder" in (child.get("class") or []) or "math-placeholder" in child.get("class", "")):
                from docx.oxml import parse_xml
                xml_str = child.get("data-xml", "")
                if xml_str:
                    try:
                        omath_el = parse_xml(xml_str.encode('utf-8'))
                        paragraph._p.append(omath_el)
                    except Exception:
                        paragraph.add_run(child.get_text())
                else:
                    paragraph.add_run(child.get_text())
            elif tag == "span" and ("docx-footnote" in (child.get("class") or []) or "docx-footnote" in child.get("class", "")):
                f_id = child.get("data-footnote-id")
                f_text = child.get("data-footnote-text", "")
                if not f_id:
                    # Generate a ID based on footnotes count + unique document tag position
                    import random
                    f_id = str(random.randint(100, 99999))
                
                try:
                    from docx.oxml import parse_xml, OxmlElement
                    from docx.oxml.ns import qn
                    from lxml import etree
                    
                    footnotes_part = _get_or_create_footnotes_part(docx_doc)
                    footnotes_el = parse_xml(footnotes_part.blob)
                    
                    # Check if this footnote ID is already in footnotes xml
                    existing = footnotes_el.find(f'.//w:footnote[@w:id="{f_id}"]', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if existing is None:
                        new_footnote = OxmlElement('w:footnote')
                        new_footnote.set(qn('w:id'), f_id)
                        
                        p_el = OxmlElement('w:p')
                        
                        r_el = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        rPr.append(OxmlElement('w:footnoteRef'))
                        r_el.append(rPr)
                        p_el.append(r_el)
                        
                        r_num_el = OxmlElement('w:r')
                        t_num_el = OxmlElement('w:t')
                        num_text = child.get_text().strip().strip("[]")
                        t_num_el.text = f"{num_text} " if num_text else f"{f_id} "
                        r_num_el.append(t_num_el)
                        p_el.append(r_num_el)
                        
                        r_text_el = OxmlElement('w:r')
                        t_text_el = OxmlElement('w:t')
                        t_text_el.text = f_text
                        r_text_el.append(t_text_el)
                        p_el.append(r_text_el)
                        
                        new_footnote.append(p_el)
                        footnotes_el.append(new_footnote)
                        footnotes_part._blob = etree.tostring(footnotes_el, encoding='utf-8', xml_declaration=True)
                    
                    # Add footnoteReference in main document paragraph run
                    run = paragraph.add_run()
                    ft_ref = OxmlElement('w:footnoteReference')
                    ft_ref.set(qn('w:id'), f_id)
                    run._r.append(ft_ref)
                except Exception:
                    run = paragraph.add_run(f"[{child.get_text()}]")
                    run.font.superscript = True
            elif tag == "span":
                # General span font style parser
                style_str = child.get("style", "")
                child_font_name = font_name
                child_font_size = font_size
                
                if style_str:
                    import re
                    m_fam = re.search(r"font-family:\s*([^;]+)", style_str)
                    if m_fam:
                        child_font_name = m_fam.group(1).strip().strip("'\"")
                    m_size = re.search(r"font-size:\s*([\d.]+)\s*pt", style_str)
                    if m_size:
                        try:
                            child_font_size = float(m_size.group(1))
                        except Exception:
                            pass
                _parse_inline_elements(child, paragraph, docx_doc, bold=bold, italic=italic, underline=underline, strike=strike, font_name=child_font_name, font_size=child_font_size)
            else:
                _parse_inline_elements(child, paragraph, docx_doc, bold=bold, italic=italic, underline=underline, strike=strike)


def _apply_paragraph_style(p, style_str):
    if not style_str:
        return
    import re
    from docx.shared import Inches, Pt
    if "text-align: center" in style_str:
        p.alignment = 1
    elif "text-align: right" in style_str:
        p.alignment = 2
    elif "text-align: justify" in style_str:
        p.alignment = 3
        
    m_left = re.search(r"margin-left:\s*(-?[\d.]+)\s*in", style_str)
    if m_left:
        try:
            p.paragraph_format.left_indent = Inches(float(m_left.group(1)))
        except Exception:
            pass
            
    m_right = re.search(r"margin-right:\s*(-?[\d.]+)\s*in", style_str)
    if m_right:
        try:
            p.paragraph_format.right_indent = Inches(float(m_right.group(1)))
        except Exception:
            pass
            
    m_bottom = re.search(r"margin-bottom:\s*(-?[\d.]+)\s*pt", style_str)
    if m_bottom:
        try:
            p.paragraph_format.space_after = Pt(float(m_bottom.group(1)))
        except Exception:
            pass
            
    m_top = re.search(r"margin-top:\s*(-?[\d.]+)\s*pt", style_str)
    if m_top:
        try:
            p.paragraph_format.space_before = Pt(float(m_top.group(1)))
        except Exception:
            pass
            
    m_lh = re.search(r"line-height:\s*([\d.]+)\s*(pt|in|cm)?", style_str)
    if m_lh:
        try:
            val = float(m_lh.group(1))
            unit = m_lh.group(2)
            if not unit:
                p.paragraph_format.line_spacing = val
            elif unit == "pt":
                p.paragraph_format.line_spacing = Pt(val)
        except Exception:
            pass


def _parse_soup_nodes(container_el, docx_doc) -> None:
    for el in container_el.children:
        if isinstance(el, NavigableString):
            if el.strip():
                docx_doc.add_paragraph(el.strip())
            continue

        tag = el.name

        if tag == "div" and ("docx-column-section" in (el.get("class") or []) or "docx-column-section" in el.get("class", "")):
            style_str = el.get("style", "")
            import re
            m = re.search(r"column-count:\s*(\d+)", style_str)
            num_cols = int(m.group(1)) if m else 1
            
            # Start continuous section for columns
            sec = docx_doc.add_section(0) # 0 = CONTINUOUS
            sec.top_margin = Inches(1)
            sec.bottom_margin = Inches(1)
            sec.left_margin = Inches(1)
            sec.right_margin = Inches(1)
            
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            cols = sec._sectPr.find(qn('w:cols'))
            if cols is None:
                cols = OxmlElement('w:cols')
                sec._sectPr.append(cols)
            cols.set(qn('w:num'), str(num_cols))
            
            _parse_soup_nodes(el, docx_doc)
            
            # Revert back to 1-column continuous section
            end_sec = docx_doc.add_section(0)
            end_sec.top_margin = Inches(1)
            end_sec.bottom_margin = Inches(1)
            end_sec.left_margin = Inches(1)
            end_sec.right_margin = Inches(1)
            
            end_cols = end_sec._sectPr.find(qn('w:cols'))
            if end_cols is None:
                end_cols = OxmlElement('w:cols')
                end_sec._sectPr.append(end_cols)
            end_cols.set(qn('w:num'), '1')
            continue

        # Check if this element is a table or wraps a table (e.g. ocr-detected-table-wrap)
        table_el = None
        if tag == "table":
            table_el = el
        elif tag == "div" and el.find("table"):
            table_el = el.find("table")

        if table_el:
            rows = table_el.find_all("tr")
            if not rows:
                continue

            max_cols = 0
            for row in rows:
                cols = row.find_all(["td", "th"])
                max_cols = max(max_cols, len(cols))

            if max_cols == 0:
                continue

            table = docx_doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'

            for row_idx, row in enumerate(rows):
                cols = row.find_all(["td", "th"])
                for col_idx, cell in enumerate(cols):
                    if col_idx >= max_cols:
                        continue
                    
                    colspan = int(cell.get("colspan", 1))
                    rowspan = int(cell.get("rowspan", 1))
                    
                    docx_cell = table.cell(row_idx, col_idx)
                    
                    if colspan > 1 or rowspan > 1:
                        try:
                            target_row = min(row_idx + rowspan - 1, len(rows) - 1)
                            target_col = min(col_idx + colspan - 1, max_cols - 1)
                            target_cell = table.cell(target_row, target_col)
                            docx_cell.merge(target_cell)
                        except Exception:
                            # Ignore merge errors due to complex/non-rectangular HTML table structures
                            pass

                    if docx_cell.paragraphs:
                        p = docx_cell.paragraphs[0]
                    else:
                        p = docx_cell.add_paragraph()
                    _parse_inline_elements(cell, p, docx_doc)
            continue

        if tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag[1])
            heading_text = el.get_text().strip()
            if heading_text:
                p = docx_doc.add_heading(heading_text, level=level)
                style = el.get("style", "")
                _apply_paragraph_style(p, style)

        elif tag in ["ul", "ol"]:
            style = 'List Bullet' if tag == "ul" else 'List Number'
            for li in el.find_all("li"):
                p = docx_doc.add_paragraph(style=style)
                li_style = li.get("style", "")
                _apply_paragraph_style(p, li_style)
                _parse_inline_elements(li, p, docx_doc)

        elif tag == "p":
            p = docx_doc.add_paragraph()
            style = el.get("style", "")
            _apply_paragraph_style(p, style)
            _parse_inline_elements(el, p, docx_doc)

        else:
            if tag == "img":
                p = docx_doc.add_paragraph()
                _add_image_to_paragraph(el, p)
            else:
                p = docx_doc.add_paragraph()
                style = el.get("style", "")
                _apply_paragraph_style(p, style)
                _parse_inline_elements(el, p, docx_doc)


def parse_html_to_docx(html_content: str, docx_doc) -> None:
    if not html_content:
        return

    soup = BeautifulSoup(html_content, "html.parser")
    _parse_soup_nodes(soup, docx_doc)


def documents_to_docx(pages) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for i, page in enumerate(pages):
        if i > 0:
            doc.add_page_break()

        if not page.revisions:
            doc.add_paragraph("[No content available for this page]")
            continue

        rev = page.revisions[-1]
        if getattr(rev, "document", None):
            doc_obj = document_for_revision(rev, page)
            html_content = doc_obj.to_html(replica=False)
            parse_html_to_docx(html_content, doc)
        else:
            lines = (rev.content or "").strip().split("\n\n")
            for paragraph_text in lines:
                if paragraph_text.strip():
                    p = doc.add_paragraph()
                    wrapped_text = auto_wrap_math(paragraph_text.strip())
                    _add_text_and_math_to_paragraph(wrapped_text, p)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()


def _inline_css_in_html(soup) -> None:
    link_tag = soup.find("link", rel="stylesheet")
    if link_tag:
        from flask import current_app
        try:
            static_folder = current_app.static_folder
            css_path = os.path.join(static_folder, "gen", "style.css")
            if not os.path.exists(css_path):
                css_path = os.path.join(static_folder, "css", "style.css")

            if os.path.exists(css_path):
                with open(css_path, encoding="utf-8") as f:
                    css_content = f.read()

                style_tag = soup.new_tag("style")
                style_tag.string = css_content
                link_tag.replace_with(style_tag)
        except Exception:
            pass


def documents_to_html_zip(project, pages, *, replica: bool = True) -> bytes:
    raw_html = documents_to_html(pages, replica=replica)
    soup = BeautifulSoup(raw_html, "html.parser")

    _inline_css_in_html(soup)

    image_keys_to_fetch = {}  # maps zip_path -> storage_key

    for img in soup.find_all("img"):
        src = img.get("src")
        if not src:
            continue

        match = re.search(r"static/uploads/([^/]+)/images/([^/?]+)", src)
        if not match:
            match = re.search(r"static/uploads/([^/]+)/([^/?]+)", src)

        if match:
            project_slug = match.group(1)
            filename = match.group(2)
            key = editor_image_key(project_slug, filename)

            local_path = f"images/{filename}"
            img["src"] = local_path
            image_keys_to_fetch[local_path] = key

    zip_buffer = io.BytesIO()
    storage = get_storage()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr("index.html", str(soup))

        for zip_path, key in image_keys_to_fetch.items():
            try:
                if storage.exists(key):
                    img_bytes = storage.read_bytes(key)
                    zip_file.writestr(zip_path, img_bytes)
            except Exception:
                pass

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def _crop_figure_image(img_path, bbox) -> bytes | None:
    """Crop only the figure portion of the scanned page image using Pillow."""
    from PIL import Image
    import io
    try:
        with Image.open(img_path) as img:
            w, h = img.size
            x1 = max(0, min(bbox[0], w))
            y1 = max(0, min(bbox[1], h))
            x2 = max(0, min(bbox[2], w))
            y2 = max(0, min(bbox[3], h))
            if x2 <= x1 or y2 <= y1:
                return None
            cropped = img.crop((x1, y1, x2, y2))
            img_byte_arr = io.BytesIO()
            fmt = img.format or 'JPEG'
            cropped.save(img_byte_arr, format=fmt)
            return img_byte_arr.getvalue()
    except Exception:
        return None


def _insert_styled_text(pdf_page, rect, html_content, fontname, fontsize):
    """Insert text into a textbox preserving bold/italic if insert_htmlbox is available."""
    from bs4 import BeautifulSoup
    if hasattr(pdf_page, "insert_htmlbox"):
        try:
            css = f"""
            body {{
                font-family: "{fontname}", sans-serif;
                font-size: {fontsize}pt;
                line-height: 1.3;
                margin: 0;
                padding: 0;
            }}
            h1 {{ font-size: 1.4em; font-weight: bold; margin: 0; }}
            h2 {{ font-size: 1.2em; font-weight: bold; margin: 0; }}
            h3 {{ font-size: 1.1em; font-weight: bold; margin: 0; }}
            p {{ margin: 0; }}
            b, strong {{ font-weight: bold; }}
            i, em {{ font-style: italic; }}
            """
            pdf_page.insert_htmlbox(rect, html_content, css=css)
            return
        except Exception:
            pass

    # Fallback to plain text
    text = BeautifulSoup(html_content, "html.parser").get_text().strip()
    pdf_page.insert_textbox(rect, text, fontname=fontname, fontsize=fontsize, align=0)


def _insert_block_image(pdf_page, rect, html_content) -> bool:
    """Check if the block content contains an <img> tag, and render it inside the PDF."""
    from bs4 import BeautifulSoup
    import re
    from kalanjiyam.utils.storage import get_storage, editor_image_key

    soup = BeautifulSoup(html_content, "html.parser")
    img_tag = soup.find("img")
    if not img_tag:
        return False

    src = img_tag.get("src")
    if not src:
        return False

    # Extract project_slug and filename
    match = re.search(r"static/uploads/([^/]+)/images/([^/?]+)", src)
    if not match:
        match = re.search(r"static/uploads/([^/]+)/([^/?]+)", src)

    if match:
        project_slug = match.group(1)
        filename = match.group(2)
        key = editor_image_key(project_slug, filename)

        storage = get_storage()
        try:
            if storage.exists(key):
                img_bytes = storage.read_bytes(key)
                pdf_page.insert_image(rect, stream=img_bytes)
                return True
        except Exception:
            pass

    # Draw fallback placeholder box
    shape = pdf_page.new_shape()
    shape.draw_rect(rect)
    shape.finish(color=(0.8, 0.8, 0.8), fill=(0.95, 0.95, 0.95), width=1)
    shape.commit()

    pdf_page.insert_textbox(
        rect,
        "[Image]",
        fontname="helv",
        fontsize=9,
        align=1
    )
    return True


def documents_to_pdf(project, pages) -> bytes:
    """Download pages compiled into a single PDF document in replica layout."""
    import fitz
    from kalanjiyam.utils.assets import get_page_image_filepath
    from kalanjiyam.utils.page_document import document_for_revision
    import os
    import re
    import glob
    from bs4 import BeautifulSoup

    doc = fitz.open()

    def _has_devanagari(text: str) -> bool:
        return bool(re.search(r"[\u0900-\u097f]", text))

    def _has_tamil(text: str) -> bool:
        return bool(re.search(r"[\u0b80-\u0bff]", text))

    # Paths to specific Noto Sans fonts for Unicode script rendering
    font_paths = {
        "devanagari": "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
        "tamil": "/usr/share/fonts/truetype/noto/NotoSansTamil-Regular.ttf",
        "latin": "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "fallback_dejavu": "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    }

    for page in pages:
        # Enforce strict A4 format
        width, height = 595, 842

        # Create clean white page
        pdf_page = doc.new_page(width=width, height=height)

        # Register Unicode fonts available on the system for this page
        registered_fonts = {}
        for font_key, path in font_paths.items():
            if os.path.exists(path):
                try:
                    pdf_page.insert_font(fontname=font_key, fontfile=path)
                    registered_fonts[font_key] = font_key
                except Exception:
                    pass

        # If no custom fonts were registered, look for any system TTF font as fallback
        if not registered_fonts:
            fallback_path = None
            for p in glob.glob("/usr/share/fonts/truetype/**/*.ttf", recursive=True):
                fallback_path = p
                break
            if fallback_path:
                try:
                    pdf_page.insert_font(fontname="general_fallback", fontfile=fallback_path)
                    registered_fonts["general_fallback"] = "general_fallback"
                except Exception:
                    pass

        # Helper to select the best registered font name for a given block text
        def _get_font_for_text(text: str) -> str:
            if _has_devanagari(text) and "devanagari" in registered_fonts:
                return "devanagari"
            elif _has_tamil(text) and "tamil" in registered_fonts:
                return "tamil"
            elif "latin" in registered_fonts:
                return "latin"
            elif "fallback_dejavu" in registered_fonts:
                return "fallback_dejavu"
            elif "general_fallback" in registered_fonts:
                return "general_fallback"
            return "helv"

        # Render document content (text blocks, figures, tables) at coordinates
        if page.revisions:
            rev = page.revisions[-1]
            if getattr(rev, "document", None):
                doc_obj = document_for_revision(rev, page)

                # Determine if we have a scanned image to extract figures from
                img_path = None
                try:
                    img_path = get_page_image_filepath(project.slug, page.slug)
                    if img_path and not img_path.exists():
                        img_path = None
                except Exception:
                    pass

                for block in doc_obj.blocks:
                    bbox = block.bbox
                    if not bbox or len(bbox) != 4:
                        continue

                    # Scale block coordinates relative to A4 page size
                    pw = page.page_width or width
                    ph = page.page_height or height
                    scale_x = width / pw
                    scale_y = height / ph

                    rect = fitz.Rect(
                        bbox[0] * scale_x,
                        bbox[1] * scale_y,
                        bbox[2] * scale_x,
                        bbox[3] * scale_y
                    )

                    # 1. First check if block contains an inline <img> tag (uploaded image)
                    if _insert_block_image(pdf_page, rect, block.content or ""):
                        continue

                    # 2. Otherwise render other types
                    if block.type == "figure":
                        # Attempt to crop the figure from the scanned book page image
                        img_bytes = None
                        if img_path:
                            img_bytes = _crop_figure_image(img_path, bbox)
                        
                        if img_bytes:
                            try:
                                pdf_page.insert_image(rect, stream=img_bytes)
                            except Exception:
                                img_bytes = None
                                
                        if not img_bytes:
                            # Fallback: Draw placeholder box for image
                            shape = pdf_page.new_shape()
                            shape.draw_rect(rect)
                            shape.finish(color=(0.8, 0.8, 0.8), fill=(0.95, 0.95, 0.95), width=1)
                            shape.commit()

                            # Draw [Image] label in the center
                            pdf_page.insert_textbox(
                                rect,
                                "[Image]",
                                fontname="helv",
                                fontsize=9,
                                align=1
                            )
                    elif block.type == "table":
                        # Render tables with a border
                        shape = pdf_page.new_shape()
                        shape.draw_rect(rect)
                        shape.finish(color=(0.7, 0.7, 0.7), width=1)
                        shape.commit()

                        text = (block.content or "").strip()
                        active_font = _get_font_for_text(text)
                        # Render table content (using small monospace/fallback font)
                        _insert_styled_text(pdf_page, rect + (4, 4, -4, -4), text, active_font, 8)
                    else:
                        text = (block.content or "").strip()
                        if not text:
                            continue

                        # Select active font for this text block
                        active_font = _get_font_for_text(text)

                        # Insert the text block preserving styles
                        _insert_styled_text(pdf_page, rect, text, active_font, 10)
            else:
                # Fallback: if there's no structured document model, draw raw content
                text = rev.content or ""
                margin_rect = pdf_page.rect + (36, 36, -36, -36)
                active_font = _get_font_for_text(text)
                _insert_styled_text(pdf_page, margin_rect, text, active_font, 12)

    pdf_bytes = doc.write()
    doc.close()
    return pdf_bytes

