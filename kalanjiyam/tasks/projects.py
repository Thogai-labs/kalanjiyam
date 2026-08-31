"""Background tasks for proofing projects."""

import logging
import tempfile
from pathlib import Path

# NOTE: `fitz` is the internal package name for PyMuPDF. PyPI hosts another
# package called `fitz` (https://pypi.org/project/fitz/) that is completely
# unrelated to PDF parsing.
import fitz
from PIL import Image, ImageOps
from slugify import slugify

from kalanjiyam import database as db
from kalanjiyam import queries as q
from kalanjiyam.tasks import app
from kalanjiyam.tasks.utils import CeleryTaskStatus, TaskStatus
from kalanjiyam.utils.quotas import add_storage_usage_for_project
from kalanjiyam.utils.storage import Storage, get_storage, page_image_key
from config import create_config_only_app


def _split_pdf_into_pages(
    pdf_path: Path, slug: str, storage: Storage, task_status: TaskStatus, org_slug: str = "open-tenant"
) -> int:
    """Split the given PDF into N .jpg images, one image per page.

    Each page image is saved to `storage` as it is rendered.

    :param pdf_path: local filesystem path to the PDF we should process.
    :param slug: the project slug, which determines the storage keys.
    :param storage: the storage backend to save page images to.
    :return: the page count, which we use downstream.
    """
    doc = fitz.open(pdf_path)
    task_status.progress(0, doc.page_count, doc_type="pdf")
    with tempfile.TemporaryDirectory() as tmp_dir:
        for page in doc:
            n = page.number + 1
            pix = page.get_pixmap(dpi=200)
            tmp_path = Path(tmp_dir) / f"{n}.jpg"
            pix.pil_save(tmp_path, optimize=True)
            storage.save(page_image_key(slug, str(n), org_slug=org_slug), tmp_path)
            tmp_path.unlink()
            task_status.progress(n, doc.page_count, doc_type="pdf")
    return doc.page_count


def process_page_image_for_storage(im: Image.Image) -> Image.Image:
    """Normalize a page scan image to 200 DPI RGB.

    1. Transposes orientation according to EXIF metadata (camera/phone rotation).
    2. Converts color space to RGB.
    3. Resamples to 200 DPI if scanned at higher resolution (> 200 DPI)
       or if uncalibrated image dimensions exceed standard 200 DPI document sizes.
    """
    im = ImageOps.exif_transpose(im)
    if im.mode != "RGB":
        im = im.convert("RGB")

    dpi_info = im.info.get("dpi")
    src_dpi = None
    if isinstance(dpi_info, (tuple, list)) and len(dpi_info) >= 1:
        try:
            val = float(dpi_info[0])
            if len(dpi_info) >= 2:
                val_y = float(dpi_info[1])
                if val_y > 0 and val > 0:
                    val = max(val, val_y)
            if val > 0:
                src_dpi = val
        except (ValueError, TypeError):
            pass
    elif isinstance(dpi_info, (int, float)) and dpi_info > 0:
        src_dpi = float(dpi_info)

    w, h = im.size
    scale = 1.0

    # If scanner embedded a calibrated DPI > 200 (e.g. 300, 400, 600 DPI)
    if src_dpi and src_dpi > 200:
        scale = 200.0 / src_dpi
    elif (not src_dpi or src_dpi <= 100) and max(w, h) > 2400:
        # Camera / phone scan or uncalibrated high-res scan:
        # Scale longest edge to ~2400 px (standard document height at 200 DPI)
        scale = 2400.0 / max(w, h)

    if scale < 1.0:
        new_w = max(1, int(round(w * scale)))
        new_h = max(1, int(round(h * scale)))
        im = im.resize((new_w, new_h), Image.Resampling.LANCZOS)

    return im


def _add_project_to_database(
    display_title: str,
    slug: str,
    num_pages: int,
    creator_id: int | None,
    require_org: bool,
    fingerprint_id: str | None = None,
    org_slug: str = "open-tenant",
):
    """Create a project on the database.

    :param display_title: the project title
    :param num_pages: the number of pages in the project
    """

    logging.info(f"Creating project (slug = {slug}) ...")
    session = q.get_session()
    board = db.Board(title=f"{slug} discussion board")
    session.add(board)
    session.flush()

    project = db.Project(
        slug=slug,
        display_title=display_title,
        creator_id=creator_id,
        fingerprint_id=fingerprint_id,
    )
    project.board_id = board.id
    session.add(project)
    session.flush()

    logging.info(f"Fetching project and status (slug = {slug}) ...")
    unreviewed = session.query(db.PageStatus).filter_by(name="reviewed-0").one()

    logging.info(f"Creating {num_pages} Page entries (slug = {slug}) ...")
    for n in range(1, num_pages + 1):
        session.add(
            db.Page(
                project_id=project.id,
                slug=str(n),
                order=n,
                status_id=unreviewed.id,
            )
        )
    creator = session.query(db.User).filter_by(id=creator_id).first() if creator_id else None
    # Auto-assign projects to the target or creator's organization for tenant isolation.
    target_group = None
    if org_slug and org_slug != "open-tenant":
        target_group = session.query(db.Group).filter_by(slug=org_slug).first()
    if not target_group and creator:
        from kalanjiyam.utils.org_access import user_organization_id
        creator_org_id = user_organization_id(creator)
        if creator_org_id:
            target_group = session.query(db.Group).filter_by(id=creator_org_id).first()

    if target_group:
        session.add(db.ProjectGroups(group_id=target_group.id, project_id=project.id))
    elif not creator and fingerprint_id:
        # Guests default to the open-tenant workspace
        try:
            open_tenant = q.get_or_create_open_tenant()
            session.add(db.ProjectGroups(group_id=open_tenant.id, project_id=project.id))
        except Exception:
            pass
    elif creator and require_org:
        raise ValueError("Project creator must belong to an organization.")
    session.commit()


def _extract_docx_images(doc, project_slug, storage, org_slug: str = "open-tenant") -> dict:
    image_mapping = {}
    for r_id, rel in doc.part.rels.items():
        if "image" in rel.reltype or "image" in rel.target_ref:
            try:
                img_bytes = rel.target_part.blob
                ext = rel.target_ref.split(".")[-1]
                filename = f"image_{r_id}.{ext}"
                from kalanjiyam.utils.storage import editor_image_key
                key = editor_image_key(project_slug, filename, org_slug=org_slug)
                storage.save(key, img_bytes)
                image_mapping[r_id] = filename
            except Exception as ex:
                logging.warning(f"Failed to extract image relation {r_id}: {ex}")
    return image_mapping


def _parse_run_to_html(run, child, project_slug, image_mapping) -> str:
    import html
    from docx.oxml.ns import qn
    
    text = run.text or ""
    if not text:
        embed_id = None
        for blip in child.xpath('.//*[local-name()="blip"]'):
            embed_id = blip.get(qn('r:embed')) or blip.get(qn('r:link'))
            if embed_id:
                break
        if not embed_id:
            for img_data in child.xpath('.//*[local-name()="imagedata"]'):
                embed_id = img_data.get(qn('r:id')) or img_data.get(qn('r:href'))
                if embed_id:
                    break
                    
        if embed_id and embed_id in image_mapping:
            filename = image_mapping[embed_id]
            width_attr = ""
            height_attr = ""
            extents = child.xpath('.//*[local-name()="extent"]')
            cx = cy = None
            if extents:
                cx = extents[0].get('cx')
                cy = extents[0].get('cy')
            if not cx or not cy:
                exts = child.xpath('.//*[local-name()="ext"]')
                if exts:
                    cx = exts[0].get('cx')
                    cy = exts[0].get('cy')
            
            style_dims = ""
            if cx and cy:
                try:
                    width_in = int(cx) / 914400.0
                    height_in = int(cy) / 914400.0
                    style_dims = f' style="width: {width_in:.4f}in; height: {height_in:.4f}in;"'
                except Exception:
                    pass
            return f'<img src="/static/uploads/{project_slug}/images/{filename}" class="inline-block align-middle max-h-16 mx-1" alt="Image"{style_dims} />'
            
        if child.xpath('.//*[local-name()="br"]'):
            return '<br/>'
        return ""
        
    run_html = html.escape(text)
    styles = []
    if run.font:
        font_name = run.font.name
        if not font_name:
            rPr = child.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    font_name = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi')) or rFonts.get(qn('w:cs'))
        if font_name:
            styles.append(f"font-family: {font_name};")
        if run.font.size:
            styles.append(f"font-size: {run.font.size.pt}pt;")
    style_str = " ".join(styles)
    if style_str:
        run_html = f'<span style="{style_str}">{run_html}</span>'
        
    if run.bold:
        run_html = f'<strong>{run_html}</strong>'
    if run.italic:
        run_html = f'<em>{run_html}</em>'
    if run.underline:
        run_html = f'<u>{run_html}</u>'
    if run.font and run.font.strike:
        run_html = f'<s>{run_html}</s>'
    return run_html


def _parse_paragraph_to_html(p, project_slug, image_mapping, footnotes=None) -> str | tuple:
    import html
    from docx.oxml.ns import qn
    from docx.text.run import Run
    from lxml import etree

    html_runs = []
    
    for child in p._p.iterchildren():
        tag = child.tag
        if tag.endswith('pPr'):
            continue
            
        elif tag.endswith('r'):
            # Check for footnote reference
            ft_ref = child.find(qn('w:footnoteReference'))
            if ft_ref is not None:
                f_id = ft_ref.get(qn('w:id'))
                if f_id and footnotes and f_id in footnotes:
                    text_content = f"[{f_id}]"
                    html_runs.append(f'<span class="docx-footnote" data-footnote-id="{f_id}" data-footnote-text="{html.escape(footnotes[f_id])}">{text_content}</span>')
                    continue
            
            run = Run(child, p)
            run_html = _parse_run_to_html(run, child, project_slug, image_mapping)
            if run_html:
                html_runs.append(run_html)
                
        elif tag.endswith('hyperlink'):
            r_id = child.get(qn('r:id'))
            href = ""
            if r_id and r_id in p.part.rels:
                href = p.part.rels[r_id].target_ref
            
            link_runs = []
            for sub_child in child.iterchildren():
                if sub_child.tag.endswith('r'):
                    run = Run(sub_child, p)
                    run_html = _parse_run_to_html(run, sub_child, project_slug, image_mapping)
                    if run_html:
                        link_runs.append(run_html)
            
            link_content = "".join(link_runs)
            html_runs.append(f'<a href="{html.escape(href)}">{link_content}</a>')
            
        elif tag.endswith('oMath') or tag.endswith('oMathPara'):
            oMath_xml = etree.tostring(child, encoding='utf-8').decode('utf-8')
            html_runs.append(f'<span class="math-placeholder" data-xml="{html.escape(oMath_xml)}">[Equation]</span>')

    content = "".join(html_runs)
    if not content.strip():
        content = "&nbsp;"

    style_attrs = []
    if p.alignment is not None:
        align_val = p.alignment
        if align_val == 1:
            style_attrs.append("text-align: center;")
        elif align_val == 2:
            style_attrs.append("text-align: right;")
        elif align_val == 3:
            style_attrs.append("text-align: justify;")
            
    # Spacing and Indentations
    try:
        if p.paragraph_format:
            pf = p.paragraph_format
            if pf.left_indent is not None:
                style_attrs.append(f"margin-left: {pf.left_indent.inches:.4f}in;")
            if pf.right_indent is not None:
                style_attrs.append(f"margin-right: {pf.right_indent.inches:.4f}in;")
            if pf.space_after is not None:
                style_attrs.append(f"margin-bottom: {pf.space_after.pt:.2f}pt;")
            if pf.space_before is not None:
                style_attrs.append(f"margin-top: {pf.space_before.pt:.2f}pt;")
            if pf.line_spacing is not None:
                if isinstance(pf.line_spacing, float):
                    style_attrs.append(f"line-height: {pf.line_spacing};")
                else:
                    style_attrs.append(f"line-height: {pf.line_spacing.pt:.2f}pt;")
    except Exception:
        pass

    style_str = " ".join(style_attrs)

    style_name = (p.style.name or "").lower() if p.style else ""
    numPr = p._p.pPr.find(qn('w:numPr')) if (p._p.pPr is not None) else None
    
    is_list = numPr is not None or "list" in style_name or style_name.startswith("bullet") or style_name.startswith("numbered")
    if is_list:
        is_bullet = "bullet" in style_name or "number" not in style_name
        list_type = "bullet" if is_bullet else "ordered"
        return ("li", list_type, content, style_str)

    if style_name.startswith("heading") or "heading" in style_name:
        try:
            import re
            m = re.search(r'\d+', style_name)
            level = int(m.group(0)) if m else 1
            if 1 <= level <= 6:
                tag = f"h{level}"
            else:
                tag = "h1"
        except Exception:
            tag = "h1"
        tag_style = f' style="{style_str}"' if style_str else ""
        return f'<{tag}{tag_style}>{content}</{tag}>'

    tag_style = f' style="{style_str}"' if style_str else ""
    return f'<p{tag_style}>{content}</p>'


def _parse_table_to_html(table, project_slug, image_mapping, footnotes=None) -> str:
    from docx.oxml.ns import qn
    
    rows = table.rows
    if not rows:
        return ""
        
    num_rows = len(rows)
    num_cols = max(len(row.cells) for row in rows) if num_rows > 0 else 0
    
    html_rows = []
    
    for r_idx, row in enumerate(rows):
        html_cells = []
        for c_idx, cell in enumerate(row.cells):
            coord = (r_idx, c_idx)
            cell_element = cell._tc
            
            first_coord = None
            for tr_idx, trow in enumerate(rows):
                for tc_idx, tcell in enumerate(trow.cells):
                    if tcell._tc is cell_element:
                        first_coord = (tr_idx, tc_idx)
                        break
                if first_coord:
                    break
            
            if first_coord != coord:
                continue
                
            colspan = 1
            tcPr = cell_element.get_or_add_tcPr()
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                val = gridSpan.get(qn('w:val'))
                if val:
                    try:
                        colspan = int(val)
                    except Exception:
                        pass
                    
            rowspan = 1
            next_r = r_idx + 1
            while next_r < num_rows:
                if c_idx < len(rows[next_r].cells) and rows[next_r].cells[c_idx]._tc is cell_element:
                    rowspan += 1
                    next_r += 1
                else:
                    break
            
            cell_html = []
            for p in cell.paragraphs:
                res = _parse_paragraph_to_html(p, project_slug, image_mapping, footnotes)
                if isinstance(res, tuple):
                    cell_style = f' style="{res[3]}"' if res[3] else ""
                    cell_html.append(f'<p{cell_style}>{res[2]}</p>')
                else:
                    cell_html.append(res)
            
            td_attrs = []
            if colspan > 1:
                td_attrs.append(f'colspan="{colspan}"')
            if rowspan > 1:
                td_attrs.append(f'rowspan="{rowspan}"')
                
            td_attr_str = " ".join(td_attrs)
            td_tag = f'<td {td_attr_str}>' if td_attr_str else '<td>'
            
            html_cells.append(f'{td_tag}{"".join(cell_html)}</td>')
            
        html_rows.append(f'<tr>{"".join(html_cells)}</tr>')
        
    return f'<table class="border-collapse border border-teal-300 w-full">{"".join(html_rows)}</table>'


def _compile_elements_to_html(elements) -> str:
    grouped_html = []
    list_group = []
    list_type = None

    for item in elements:
        if isinstance(item, tuple) and item[0] == "li":
            _, curr_type, content, style_attr = item
            if list_group and list_type != curr_type:
                tag = "ul" if list_type == "bullet" else "ol"
                grouped_html.append(f'<{tag}>{"".join(list_group)}</{tag}>')
                list_group = []
            list_type = curr_type
            li_style = f' style="{style_attr}"' if style_attr else ""
            list_group.append(f'<li{li_style}>{content}</li>')
        else:
            if list_group:
                tag = "ul" if list_type == "bullet" else "ol"
                grouped_html.append(f'<{tag}>{"".join(list_group)}</{tag}>')
                list_group = []
                list_type = None
            grouped_html.append(item)

    if list_group:
        tag = "ul" if list_type == "bullet" else "ol"
        grouped_html.append(f'<{tag}>{"".join(list_group)}</{tag}>')

    return "".join(grouped_html)


def _get_paragraph_plain_text(p) -> str:
    from docx.text.run import Run
    text_parts = []
    for child in p._p.iterchildren():
        tag = child.tag
        if tag.endswith('r'):
            run = Run(child, p)
            text_parts.append(run.text or "")
        elif tag.endswith('hyperlink'):
            for sub_child in child.iterchildren():
                if sub_child.tag.endswith('r'):
                    run = Run(sub_child, p)
                    text_parts.append(run.text or "")
        elif tag.endswith('oMath') or tag.endswith('oMathPara'):
            t_elements = child.xpath('.//*[local-name()="t"]')
            for t_el in t_elements:
                text_parts.append(t_el.text or "")
    return "".join(text_parts)


def _segment_docx(doc, slug, image_mapping) -> list[tuple[str, str]]:
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.ns import qn

    # Load all footnotes from docx package relations
    footnotes = {}
    try:
        from docx.oxml import parse_xml
        for rel_id, rel in doc.part.rels.items():
            if "footnotes" in rel.reltype:
                footnotes_el = parse_xml(rel.target_part.blob)
                for footnote in footnotes_el.findall('.//w:footnote', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    f_id = footnote.get(qn('w:id'))
                    if f_id:
                        texts = [t.text for t in footnote.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) if t.text]
                        footnotes[f_id] = "".join(texts)
    except Exception as e:
        import logging
        logging.warning(f"Failed to parse footnotes: {e}")

    pages = []
    current_page_elements = []  # list of (element_html, cols)
    current_page_text = []

    # Get the columns of the final section in the body
    final_cols = 1
    body_sectPr = doc.element.body.find(qn('w:sectPr'))
    if body_sectPr is not None:
        cols_el = body_sectPr.find(qn('w:cols'))
        if cols_el is not None:
            val = cols_el.get(qn('w:num'))
            if val:
                try:
                    final_cols = int(val)
                except Exception:
                    pass

    def flush_page():
        if current_page_elements:
            sections_html = []
            sec_elements = []
            sec_cols = None

            for item, cols in current_page_elements:
                if sec_cols is None:
                    sec_cols = cols
                elif sec_cols != cols:
                    sec_html = _compile_elements_to_html(sec_elements)
                    if sec_cols > 1:
                        sections_html.append(f'<div class="docx-column-section" style="column-count: {sec_cols}; column-gap: 2rem;">{sec_html}</div>')
                    else:
                        sections_html.append(sec_html)
                    sec_elements = []
                    sec_cols = cols
                sec_elements.append(item)

            if sec_elements:
                sec_html = _compile_elements_to_html(sec_elements)
                if sec_cols > 1:
                    sections_html.append(f'<div class="docx-column-section" style="column-count: {sec_cols}; column-gap: 2rem;">{sec_html}</div>')
                else:
                    sections_html.append(sec_html)

            page_html = "".join(sections_html)
            page_text = "\n".join(current_page_text)
            pages.append((page_text, page_html))
            current_page_elements.clear()
            current_page_text.clear()

    # Pre-parse sections to map children to section column counts
    body_children = list(doc.element.body.iterchildren())
    element_sections = []
    current_sec = []
    
    for child in body_children:
        if child.tag.endswith('sectPr'):
            continue
        current_sec.append(child)
        if child.tag.endswith('p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                element_sections.append((current_sec, pPr.find(qn('w:sectPr'))))
                current_sec = []
    if current_sec:
        element_sections.append((current_sec, body_sectPr))

    blocks_with_cols = []
    for sec_children, sectPr in element_sections:
        cols_num = 1
        if sectPr is not None:
            cols_el = sectPr.find(qn('w:cols'))
            if cols_el is not None:
                val = cols_el.get(qn('w:num'))
                if val:
                    try:
                        cols_num = int(val)
                    except Exception:
                        pass
        for child in sec_children:
            blocks_with_cols.append((child, cols_num))

    for child, cols in blocks_with_cols:
        if child.tag.endswith('p'):
            p = Paragraph(child, doc)
            res = _parse_paragraph_to_html(p, slug, image_mapping, footnotes)
            current_page_elements.append((res, cols))
            p_text = _get_paragraph_plain_text(p)
            current_page_text.append(p_text)
            
            p_xml = child.xml
            is_break = ('w:br' in p_xml and 'w:type="page"' in p_xml) or ('w:lastRenderedPageBreak' in p_xml)
            if is_break or len("".join(current_page_text)) > 1500:
                flush_page()
                
        elif child.tag.endswith('tbl'):
            table = Table(child, doc)
            res = _parse_table_to_html(table, slug, image_mapping, footnotes)
            current_page_elements.append((res, cols))
            table_text = " ".join(_get_paragraph_plain_text(pt) for row in table.rows for cell in row.cells for pt in cell.paragraphs)
            current_page_text.append(table_text)
            if len("".join(current_page_text)) > 1500:
                flush_page()

    flush_page()
    if not pages:
        pages.append(("", "<p></p>"))
    return pages


def create_project_inner(
    *,
    display_title: str,
    pdf_key: str | None = None,
    docx_key: str | None = None,
    image_keys: list[str] | None = None,
    app_environment: str,
    creator_id: int | None,
    fingerprint_id: str | None = None,
    task_status: TaskStatus,
    org_slug: str = "open-tenant",
):
    """Split the given PDF, DOCX, or images into pages and register the project on the database."""
    logging.info(
        f'Received upload task "{display_title}" for key {pdf_key or docx_key or (f"{len(image_keys)} images" if image_keys else "none")}.'
    )

    app = create_config_only_app(app_environment)
    with app.app_context():
        session = q.get_session()
        slug = slugify(display_title)
        project = session.query(db.Project).filter_by(slug=slug).first()

        if project:
            raise ValueError(
                f'Project "{display_title}" already exists. Please choose a different title.'
            )

        storage = get_storage()

        if docx_key:
            docx_path = storage.local_copy(docx_key)
            if not docx_path.exists():
                raise ValueError(f'Source DOCX not found in storage: "{docx_key}".')

            from docx import Document
            doc = Document(docx_path)
            image_mapping = _extract_docx_images(doc, slug, storage, org_slug=org_slug)
            pages_list = _segment_docx(doc, slug, image_mapping)
            num_pages = len(pages_list)

            require_org = bool(app.config.get("DEFAULT_PROJECT_REQUIRES_ORG", True))
            _add_project_to_database(
                display_title=display_title,
                slug=slug,
                num_pages=num_pages,
                creator_id=creator_id,
                require_org=require_org,
                fingerprint_id=fingerprint_id,
                org_slug=org_slug,
            )

            db_project = session.query(db.Project).filter_by(slug=slug).one()
            unreviewed = session.query(db.PageStatus).filter_by(name="reviewed-0").one()
            import uuid

            for idx, (page_text, page_html) in enumerate(pages_list):
                page_slug = str(idx + 1)
                db_page = session.query(db.Page).filter_by(project_id=db_project.id, slug=page_slug).one()

                # Original track
                pv_orig = db.PageVersion(page_id=db_page.id, version_key="original")
                session.add(pv_orig)
                session.flush()

                doc_dict = {
                    "content_format": "html",
                    "blocks": [{
                        "id": f"b{uuid.uuid4().hex[:8]}",
                        "type": "paragraph",
                        "bbox": [0, 0, 0, 0],
                        "content": page_html,
                        "reading_order": 1
                    }]
                }

                rev_orig = db.Revision(
                    project_id=db_project.id,
                    page_id=db_page.id,
                    page_version_id=pv_orig.id,
                    status_id=unreviewed.id,
                    content=page_text,
                    content_format="html",
                    document=doc_dict
                )
                session.add(rev_orig)

                # Editable active track
                pv_target = db.PageVersion(page_id=db_page.id, version_key="role:p1")
                session.add(pv_target)
                session.flush()

                rev_target = db.Revision(
                    project_id=db_project.id,
                    page_id=db_page.id,
                    page_version_id=pv_target.id,
                    status_id=unreviewed.id,
                    content=page_text,
                    content_format="html",
                    document=doc_dict
                )
                session.add(rev_target)
                session.flush()

                # Save structured page block document to .json.gz storage
                from kalanjiyam.utils.document_storage import save_revision_document
                save_revision_document(rev_orig, doc_dict)
                save_revision_document(rev_target, doc_dict)

            # Record source DOCX size in DB metadata & system metric log, then delete from storage
            docx_size_bytes = 0
            if docx_key and storage.exists(docx_key):
                docx_size_bytes = storage.size(docx_key)
            elif docx_path.exists():
                docx_size_bytes = docx_path.stat().st_size

            meta = db_project.extracted_metadata or {}
            if "source_file" not in meta or not isinstance(meta["source_file"], dict):
                meta["source_file"] = {}
            meta["source_file"]["size_bytes"] = docx_size_bytes
            meta["source_file"]["type"] = "docx"
            meta["source_file"]["deleted_after_extraction"] = True
            db_project.extracted_metadata = meta
            from sqlalchemy.orm.attributes import flag_modified
            flag_modified(db_project, "extracted_metadata")
            session.commit()

            from kalanjiyam.utils.metrics import record_metric
            from kalanjiyam.utils.org_access import user_organization_id
            creator_user = session.query(db.User).filter_by(id=creator_id).first() if creator_id else None
            org_id = user_organization_id(creator_user) if creator_user else None
            record_metric(
                category="project_upload",
                name="docx_extracted_and_deleted",
                user_id=creator_id,
                group_id=org_id,
                status="SUCCESS",
                details={
                    "project_slug": slug,
                    "docx_key": docx_key,
                    "source_file_size_bytes": docx_size_bytes,
                    "num_pages": num_pages,
                },
            )

            if docx_key and storage.exists(docx_key):
                storage.delete(docx_key)

            add_storage_usage_for_project(slug)
            # One index task for the finished project, not one per page.
            from kalanjiyam.tasks.search_index import enqueue_project

            enqueue_project(db_project.id)
            doc_type = "docx"
        elif image_keys:
            doc_type = "images"
            num_pages = len(image_keys)
            require_org = bool(app.config.get("DEFAULT_PROJECT_REQUIRES_ORG", True))
            _add_project_to_database(
                display_title=display_title,
                slug=slug,
                num_pages=num_pages,
                creator_id=creator_id,
                require_org=require_org,
                fingerprint_id=fingerprint_id,
                org_slug=org_slug,
            )

            total_images_size_bytes = 0
            task_status.progress(0, num_pages, doc_type="images")

            for idx, img_key in enumerate(image_keys, start=1):
                local_src = storage.local_copy(img_key)
                if not local_src.exists():
                    raise ValueError(f'Source image not found in storage: "{img_key}".')

                total_images_size_bytes += local_src.stat().st_size
                dest_key = page_image_key(slug, str(idx), org_slug=org_slug)

                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp_dest:
                    tmp_dest_path = Path(tmp_dest.name)
                    try:
                        with Image.open(local_src) as im:
                            im = process_page_image_for_storage(im)
                            im.save(tmp_dest_path, "JPEG", quality=90, optimize=True, dpi=(200, 200))
                        storage.save(dest_key, tmp_dest_path)
                    finally:
                        if tmp_dest_path.exists():
                            tmp_dest_path.unlink()

                # Clean up staged raw image from storage
                if storage.exists(img_key):
                    storage.delete(img_key)

                task_status.progress(idx, num_pages, doc_type="images")

            db_project = session.query(db.Project).filter_by(slug=slug).one()
            meta = db_project.extracted_metadata or {}
            if "source_file" not in meta or not isinstance(meta["source_file"], dict):
                meta["source_file"] = {}
            meta["source_file"]["size_bytes"] = total_images_size_bytes
            meta["source_file"]["type"] = "images"
            meta["source_file"]["num_images"] = num_pages
            meta["source_file"]["deleted_after_extraction"] = True
            db_project.extracted_metadata = meta
            from sqlalchemy.orm.attributes import flag_modified

            flag_modified(db_project, "extracted_metadata")
            session.commit()

            from kalanjiyam.utils.metrics import record_metric
            from kalanjiyam.utils.org_access import user_organization_id

            creator_user = (
                session.query(db.User).filter_by(id=creator_id).first()
                if creator_id
                else None
            )
            org_id = user_organization_id(creator_user) if creator_user else None
            record_metric(
                category="project_upload",
                name="images_extracted_and_deleted",
                user_id=creator_id,
                group_id=org_id,
                status="SUCCESS",
                details={
                    "project_slug": slug,
                    "source_file_size_bytes": total_images_size_bytes,
                    "num_pages": num_pages,
                },
            )

            add_storage_usage_for_project(slug)
        else:
            doc_type = "pdf"
            pdf_path = storage.local_copy(pdf_key)
            if not pdf_path.exists():
                raise ValueError(f'Source PDF not found in storage: "{pdf_key}".')

            # Record source PDF size before extraction and deletion
            pdf_size_bytes = 0
            if pdf_key and storage.exists(pdf_key):
                pdf_size_bytes = storage.size(pdf_key)
            elif pdf_path.exists():
                pdf_size_bytes = pdf_path.stat().st_size

            num_pages = _split_pdf_into_pages(pdf_path, slug, storage, task_status, org_slug=org_slug)
            require_org = bool(app.config.get("DEFAULT_PROJECT_REQUIRES_ORG", True))
            _add_project_to_database(
                display_title=display_title,
                slug=slug,
                num_pages=num_pages,
                creator_id=creator_id,
                require_org=require_org,
                fingerprint_id=fingerprint_id,
                org_slug=org_slug,
            )

            # Update DB project metadata and metrics log
            db_project = session.query(db.Project).filter_by(slug=slug).one()
            meta = db_project.extracted_metadata or {}
            if "source_file" not in meta or not isinstance(meta["source_file"], dict):
                meta["source_file"] = {}
            meta["source_file"]["size_bytes"] = pdf_size_bytes
            meta["source_file"]["type"] = "pdf"
            meta["source_file"]["deleted_after_extraction"] = True
            db_project.extracted_metadata = meta
            from sqlalchemy.orm.attributes import flag_modified
            flag_modified(db_project, "extracted_metadata")
            session.commit()

            from kalanjiyam.utils.metrics import record_metric
            from kalanjiyam.utils.org_access import user_organization_id
            creator_user = session.query(db.User).filter_by(id=creator_id).first() if creator_id else None
            org_id = user_organization_id(creator_user) if creator_user else None
            record_metric(
                category="project_upload",
                name="pdf_extracted_and_deleted",
                user_id=creator_id,
                group_id=org_id,
                status="SUCCESS",
                details={
                    "project_slug": slug,
                    "pdf_key": pdf_key,
                    "source_file_size_bytes": pdf_size_bytes,
                    "num_pages": num_pages,
                },
            )

            # Delete source PDF from storage
            if pdf_key and storage.exists(pdf_key):
                storage.delete(pdf_key)

            add_storage_usage_for_project(slug)

    task_status.success(num_pages, slug, doc_type=doc_type)
    return {
        "current": num_pages,
        "total": num_pages,
        "slug": slug,
        "num_pages": num_pages,
        "doc_type": doc_type,
    }


@app.task(bind=True)
def create_project(
    self,
    *,
    display_title: str,
    pdf_key: str | None = None,
    docx_key: str | None = None,
    image_keys: list[str] | None = None,
    app_environment: str,
    creator_id: int | None,
    fingerprint_id: str | None = None,
    org_slug: str = "open-tenant",
):
    """Split the given PDF, DOCX, or images into pages and register the project on the database."""
    task_status = CeleryTaskStatus(self)
    return create_project_inner(
        display_title=display_title,
        pdf_key=pdf_key,
        docx_key=docx_key,
        image_keys=image_keys,
        app_environment=app_environment,
        creator_id=creator_id,
        fingerprint_id=fingerprint_id,
        task_status=task_status,
        org_slug=org_slug,
    )


@app.task(bind=True)
def cleanup_uploaded_files_task(
    self, days: int = 7, force: bool = False, app_environment: str = "testing"
) -> int:
    """Celery task to delete uploaded source PDF and DOC/DOCX files older than `days` days.

    Only runs if AUTO_UPLOADED_FILES_CLEANUP config is enabled or `force` is True.
    """
    import os
    from flask import current_app, has_app_context
    from kalanjiyam.utils.storage import cleanup_old_uploaded_files, get_storage

    def _is_enabled(conf):
        val = conf.get("AUTO_UPLOADED_FILES_CLEANUP", False)
        if isinstance(val, bool):
            return val
        return str(val).lower() in ("true", "1", "yes")

    def _get_days(explicit_days):
        try:
            settings = q.get_system_settings()
            if settings and settings.auto_cleanup_days:
                return settings.auto_cleanup_days
        except Exception:
            pass
        return explicit_days or 7

    if has_app_context():
        enabled = _is_enabled(current_app.config)
        if not enabled and not force:
            logging.info("AUTO_UPLOADED_FILES_CLEANUP is disabled. Skipping cleanup task.")
            return 0
        target_days = _get_days(days)
        storage = get_storage()
        deleted_count = cleanup_old_uploaded_files(storage, days=target_days)
        logging.info(f"Cleaned up {deleted_count} uploaded source PDF/DOC files older than {target_days} days.")
        return deleted_count

    env = app_environment or os.getenv("KALANJIYAM_ENV", "testing")
    flask_app = create_config_only_app(env)
    with flask_app.app_context():
        enabled = _is_enabled(flask_app.config)
        if not enabled and not force:
            logging.info("AUTO_UPLOADED_FILES_CLEANUP is disabled. Skipping cleanup task.")
            return 0
        target_days = _get_days(days)
        storage = get_storage()
        deleted_count = cleanup_old_uploaded_files(storage, days=target_days)
        logging.info(f"Cleaned up {deleted_count} uploaded source PDF/DOC files older than {target_days} days.")
        return deleted_count
