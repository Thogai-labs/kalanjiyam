"""Background tasks for translation services."""

import logging
import os
from celery import group
from celery.result import GroupResult

from kalanjiyam import consts
from kalanjiyam import database as db
from kalanjiyam import queries as q
from kalanjiyam.tasks import app
from kalanjiyam.utils.translation_engine import translate_text, segment_text_for_translation
from kalanjiyam.utils.quotas import ensure_translation_quota_for_project, consume_translation_credit_for_project
from config import create_config_only_app

LOG = logging.getLogger(__name__)


def _run_translation_for_page_inner(
    app_env: str,
    project_slug: str,
    page_slug: str,
    source_lang: str = 'sa',
    target_lang: str = 'en',
    engine: str = 'google',
    revision_id: int = None,
    glossary: str = None,
) -> int:
    """Must run in the application context."""

    flask_app = create_config_only_app(app_env)
    with flask_app.app_context():
        bot_user = q.user(consts.BOT_USERNAME)
        if bot_user is None:
            raise ValueError(f'User "{consts.BOT_USERNAME}" is not defined.')

        session = q.get_session()
        project = q.project(project_slug)
        if project is None:
            raise ValueError(f'Project "{project_slug}" not found.')
        
        page = q.page(project.id, page_slug)
        if page is None:
            raise ValueError(f'Page "{page_slug}" not found in project "{project_slug}".')

        # Get the revision to translate
        if revision_id is None:
            # Use the latest revision
            if not page.revisions:
                raise ValueError(f'No revisions found for page "{page_slug}".')
            revision = page.revisions[-1]  # Latest revision
        else:
            revision = session.query(db.Revision).filter_by(id=revision_id).first()
            if not revision or revision.page_id != page.id:
                raise ValueError(f'Revision {revision_id} not found for page "{page_slug}".')

        # Check if translation already exists
        existing_translation = session.query(db.Translation).filter_by(
            page_id=page.id,
            revision_id=revision.id,
            source_language=source_lang,
            target_language=target_lang,
            translation_engine=engine
        ).first()

        if existing_translation:
            LOG.info(f"Translation already exists for {project_slug}/{page_slug} ({source_lang}->{target_lang})")
            return existing_translation.id

        ensure_translation_quota_for_project(project)

        # Check if the revision to translate is in blocks format to preserve page structure
        translated_content = ""
        document_payload = None
        content_format = "plain"
        translation_failed = False

        from kalanjiyam.utils.document_storage import load_revision_document

        rev_doc = load_revision_document(revision)
        if revision.content_format == "blocks" and rev_doc:
            import copy
            doc_data = copy.deepcopy(rev_doc)
            blocks = doc_data.get("blocks", [])
            
            if blocks:
                try:
                    from kalanjiyam.views.proofing.page import _translate_blocks
                    _translate_blocks(blocks, source_lang, target_lang, engine, glossary=glossary)
                except Exception as e:
                    LOG.error(f"Structured translation failed: {e}")
                    translation_failed = True
                
            if not translation_failed:
                from kalanjiyam.utils.page_document import PageDocument
                try:
                    translated_content = PageDocument.from_dict(doc_data).to_plain_text()
                except Exception as e:
                    LOG.error(f"Failed to generate plain text from translated document: {e}")
                    translated_content = ""
                content_format = "blocks"
                document_payload = doc_data
        else:
            # Segment text for translation
            text_segments = segment_text_for_translation(revision.content, max_length=1000)
            
            # Translate each segment
            translated_segments = []
            
            for segment in text_segments:
                if segment.strip():
                    try:
                        translation_response = translate_text(
                            segment, 
                            source_lang, 
                            target_lang, 
                            engine,
                            glossary=glossary
                        )
                        translated_segments.append(translation_response.translated_text)
                    except Exception as e:
                        LOG.error(f"Translation failed for segment: {e}")
                        translation_failed = True
                        break  # Stop translation if any segment fails
                else:
                    translated_segments.append(segment)

            if not translation_failed:
                # Combine translated segments
                translated_content = '\n\n'.join(translated_segments)
                content_format = "plain"
                document_payload = None

        # Only create translation record if translation was successful
        if not translation_failed:
            # Create translation record
            translation = db.Translation(
                page_id=page.id,
                revision_id=revision.id,
                author_id=bot_user.id,
                content=translated_content,
                source_language=source_lang,
                target_language=target_lang,
                translation_engine=engine,
                status='completed'
            )
            session.add(translation)

            # Create page version and revision following the OCR version track system
            version_key = f"translation:{engine}:{source_lang}->{target_lang}"
            pv = session.query(db.PageVersion).filter_by(
                page_id=page.id,
                version_key=version_key
            ).first()
            current_ver = pv.version if pv else 0

            from kalanjiyam.utils.revisions import add_revision
            from kalanjiyam.enums import SitePageStatus

            summary = f"Translation: {engine} {source_lang}->{target_lang}"
            add_revision(
                page=page,
                summary=summary,
                content=translated_content,
                status=SitePageStatus.R0,
                version=current_ver,
                author_id=bot_user.id,
                document=document_payload,
                content_format=content_format,
                version_key=version_key,
            )

            consume_translation_credit_for_project(project)
            session.commit()

            # Record UI Batch Translation metrics in BatchItem / BatchOcrPage
            try:
                trans_latency_ms = (time.time() - start_time) * 1000.0
                trans_data_bytes = len(translated_content.encode('utf-8'))
                
                batch_item = session.query(BatchItem).filter_by(project_id=project.id).order_by(BatchItem.id.desc()).first()
                if batch_item:
                    p_num = int(page_slug) if page_slug.isdigit() else page.order
                    ocr_page = session.query(BatchOcrPage).filter_by(batch_item_id=batch_item.id, page_number=p_num).first()
                    if not ocr_page:
                        ocr_page = BatchOcrPage(
                            batch_item_id=batch_item.id,
                            chunk_id=None,
                            page_number=p_num,
                            status='PENDING'
                        )

                    ocr_page.translation_latency_ms = trans_latency_ms
                    ocr_page.translation_data_size_bytes = trans_data_bytes
                    ocr_page.source_lang = source_lang
                    ocr_page.target_lang = target_lang
                    ocr_page.status = 'COMPLETED'
                    ocr_page.completed_at = datetime.utcnow()
                    session.add(ocr_page)

                    # Update cumulative item metrics
                    batch_item.total_translation_latency_ms = (batch_item.total_translation_latency_ms or 0) + trans_latency_ms
                    batch_item.translation_data_size_bytes = (batch_item.translation_data_size_bytes or 0) + trans_data_bytes
                    batch_item.source_lang = source_lang
                    batch_item.target_lang = target_lang
                    
                    completed_count = session.query(BatchOcrPage).filter_by(batch_item_id=batch_item.id, status='COMPLETED').count()
                    if completed_count >= (batch_item.total_pages or 1):
                        batch_item.status = 'COMPLETED'
                        batch_item.completed_at = datetime.utcnow()
                        if batch_item.job:
                            batch_item.job.status = 'COMPLETED'
                            batch_item.job.completed_at = datetime.utcnow()

                    session.commit()
            except Exception as metric_err:
                LOG.warning(f"Error recording UI batch translation metrics: {metric_err}")
            
            LOG.info(f"Translation completed for {project_slug}/{page_slug} ({source_lang}->{target_lang})")
            return translation.id
        else:
            LOG.warning(f"Translation failed for {project_slug}/{page_slug} ({source_lang}->{target_lang}) - no translation record created")
            return None


@app.task(bind=True)
def run_translation_for_page(
    self,
    *,
    app_env: str,
    project_slug: str,
    page_slug: str,
    source_lang: str = 'sa',
    target_lang: str = 'en',
    engine: str = 'google',
    revision_id: int = None,
    glossary: str = None,
):
    """Run translation for a single page."""
    try:
        return _run_translation_for_page_inner(
            app_env,
            project_slug,
            page_slug,
            source_lang,
            target_lang,
            engine,
            revision_id,
            glossary,
        )
    except Exception as e:
        LOG.error(f"Translation task failed for {project_slug}/{page_slug}: {e}")
        raise


def run_translation_for_project(
    app_env: str,
    project: db.Project,
    source_lang: str = 'sa',
    target_lang: str = 'en',
    engine: str = 'google',
    revision_id: int = None,
    queue: str | None = None,
    glossary: str = None,
) -> GroupResult | None:
    """Create a `group` task to run translation on a project.

    Usage:

    >>> r = run_translation_for_project(...)
    >>> progress = r.completed_count() / len(r.results)

    :param app_env: Application environment
    :param project: Project to run translation on
    :param source_lang: Source language code
    :param target_lang: Target language code
    :param engine: Translation engine to use
    :param revision_id: Specific revision ID to translate (optional)
    :param queue: The Celery queue name to route tasks to
    :return: the Celery result, or ``None`` if no tasks were run.
    """
    flask_app = create_config_only_app(app_env)
    with flask_app.app_context():
        from kalanjiyam.models.batch import BatchJob, BatchItem, BatchOcrPage
        session = q.get_session()
        db_project = session.query(db.Project).get(project.id)
        if not db_project:
            return None
        # Get pages that have revisions
        pages_with_revisions = [p for p in db_project.pages if p.revisions]

        if pages_with_revisions:
            # Create BatchJob and BatchItem for UI-triggered batch translation
            batch_job = BatchJob(
                target_uri=f"ui://translation/{project.slug}",
                status='IN_PROGRESS',
                job_type='UI_BATCH_TRANSLATION'
            )
            session.add(batch_job)
            session.flush()

            batch_item = BatchItem(
                job_id=batch_job.id,
                file_path=f"ui://translation/{project.slug}",
                project_id=db_project.id,
                status='IN_PROGRESS',
                total_pages=len(pages_with_revisions),
                source_lang=source_lang,
                target_lang=target_lang,
            )
            session.add(batch_item)
            session.flush()

            for p in pages_with_revisions:
                ocr_p = BatchOcrPage(
                    batch_item_id=batch_item.id,
                    chunk_id=None,
                    page_number=p.order,
                    status='PENDING',
                    source_lang=source_lang,
                    target_lang=target_lang,
                )
                session.add(ocr_p)
            session.commit()

    if pages_with_revisions:
        tasks = group(
            run_translation_for_page.s(
                app_env=app_env,
                project_slug=project.slug,
                page_slug=p.slug,
                source_lang=source_lang,
                target_lang=target_lang,
                engine=engine,
                revision_id=revision_id,
                glossary=glossary,
            )
            for p in pages_with_revisions
        )
        if queue:
            ret = tasks.apply_async(queue=queue)
        else:
            ret = tasks.apply_async()
        # Save the result so that we can poll for it later
        ret.save()
        return ret
    else:
        return None


@app.task(bind=True)
def run_translation_for_revision(
    self,
    *,
    app_env: str,
    revision_id: int,
    source_lang: str = 'sa',
    target_lang: str = 'en',
    engine: str = 'google',
    glossary: str = None,
):
    """Run translation for a specific revision across all pages in the project."""
    flask_app = create_config_only_app(app_env)
    with flask_app.app_context():
        session = q.get_session()
        revision = session.query(db.Revision).filter_by(id=revision_id).first()
        if not revision:
            raise ValueError(f'Revision {revision_id} not found.')
        
        project = revision.project
        if not project:
            raise ValueError(f'Project not found for revision {revision_id}.')
        
        # Run translation for the specific page of this revision
        return _run_translation_for_page_inner(
            app_env,
            project.slug,
            revision.page.slug,
            source_lang,
            target_lang,
            engine,
            revision_id,
            glossary,
        )


def _clear_translation_task_from_redis(task_id):
    """Clear translation task from Redis when it completes or fails."""
    try:
        import redis
        import os
        import json
        
        redis_client = redis.Redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379/0"))
        
        # Find the task key by scanning Redis keys
        for key in redis_client.scan_iter(match="translation_task:*"):
            task_info = redis_client.get(key)
            if task_info:
                task_data = json.loads(task_info)
                if task_data.get('task_id') == task_id:
                    redis_client.delete(key)
                    break
    except Exception as e:
        LOG.warning(f"Error clearing translation task from Redis: {e}")


@app.task(bind=True)
def run_docx_translation(
    self,
    *,
    app_env: str,
    docx_id: str,
    source_lang: str = 'sa',
    target_lang: str = 'en',
    engine: str = 'indictrans2',
    glossary: str = None,
    creator_id: int = None,
):
    """Run direct in-place translation for a standalone DOCX file."""
    import tempfile
    from pathlib import Path
    from docx import Document
    from kalanjiyam.utils.storage import get_storage, docx_upload_key, docx_translation_key
    from kalanjiyam.utils.translation_engine import translate_text
    from kalanjiyam.utils.quotas import (
        estimate_docx_pages,
        ensure_translation_quota_for_user,
        consume_translation_credits_for_user,
    )

    flask_app = create_config_only_app(app_env)
    with flask_app.app_context():
        storage = get_storage()
        upload_key = docx_upload_key(docx_id)
        trans_key = docx_translation_key(docx_id)

        local_path = storage.local_copy(upload_key)
        if not local_path.exists():
            raise ValueError(f"Uploaded DOCX not found in storage: {upload_key}")

        doc = Document(local_path)

        # Quota check
        creator = None
        if creator_id:
            from kalanjiyam import database as db
            from kalanjiyam import queries as q
            session = q.get_session()
            creator = session.query(db.User).filter_by(id=creator_id).first()

        estimated_pages = estimate_docx_pages(doc)
        if creator:
            ensure_translation_quota_for_user(creator, estimated_pages)

        # Save original DOCX data to database if enabled in config
        save_docx_data_enabled = (
            flask_app.config.get("SAVE_DOCX_DIRECT_TR_DATA", False)
            or os.getenv("SAVE_DOCX_DIRECT_TR_DATA", "").lower() in ("true", "1", "yes")
        )
        if save_docx_data_enabled:
            try:
                import json
                import redis
                original_filename = None
                try:
                    r_client = redis.Redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379/0"))
                    info_raw = r_client.get(f"docx_info:{docx_id}")
                    if info_raw:
                        info = json.loads(info_raw)
                        original_filename = info.get("original_filename")
                except Exception as re_err:
                    LOG.warning(f"Could not fetch original_filename from Redis for {docx_id}: {re_err}")

                from kalanjiyam.utils.docx_db_saver import save_original_docx_data_to_db
                save_original_docx_data_to_db(doc, docx_id, original_filename=original_filename, creator_id=creator_id)
            except Exception as save_err:
                LOG.error(f"Error saving original DOCX data to DB: {save_err}")


        import re

        SENTENCE_SPLIT_REGEX = re.compile(r'((?<=[.!?।॥])\s+|\n+)')

        def _is_matching_language(text: str, lang: str) -> bool:
            if not any(c.isalpha() for c in text):
                return False
            lang = lang.lower()
            if lang == 'en':
                if not re.search(r'[a-zA-Z]', text):
                    return False
                if re.search(r'[\u0900-\u0D7F\u0600-\u06FF]', text):
                    return False
                return True
            script_ranges = {
                'ta': r'[\u0B80-\u0BFF]', # Tamil
                'te': r'[\u0C00-\u0C7F]', # Telugu
                'kn': r'[\u0C80-\u0CFF]', # Kannada
                'ml': r'[\u0D00-\u0D7F]', # Malayalam
                'hi': r'[\u0900-\u097F]', # Hindi
                'sa': r'[\u0900-\u097F]', # Sanskrit
                'bn': r'[\u0980-\u09FF]', # Bengali
                'gu': r'[\u0A80-\u0AFF]', # Gujarati
                'or': r'[\u0B00-\u0B7F]', # Odia
                'pa': r'[\u0A00-\u0A7F]', # Punjabi / Gurmukhi
                'ur': r'[\u0600-\u06FF]', # Urdu / Arabic
                'mr': r'[\u0900-\u097F]', # Marathi
            }
            if lang in script_ranges:
                return bool(re.search(script_ranges[lang], text))
            return True

        def _is_safe_run(run):
            xml = run._element.xml
            if "w:drawing" in xml or "m:oMath" in xml or "m:oMathPara" in xml:
                return False
            return True

        def translate_paragraph_in_place(p, source, target, eng, glossary=None):
            text = p.text
            if not text or not text.strip():
                return
            
            try:
                # Segment text into sentences to selectively translate matching language only
                subparts = SENTENCE_SPLIT_REGEX.split(text)
                reconstructed = []
                for j in range(len(subparts)):
                    # Even indices are text segments, odd indices are delimiters
                    if j % 2 == 0:
                        sub_text = subparts[j]
                        if sub_text and sub_text.strip() and _is_matching_language(sub_text, source):
                            response = translate_text(sub_text, source, target, eng, glossary=glossary)
                            reconstructed.append(response.translated_text)
                        else:
                            reconstructed.append(sub_text)
                    else:
                        reconstructed.append(subparts[j])
                
                translated = "".join(reconstructed)
                if translated == text:
                    return # No changes made
                
                runs = p.runs
                if runs:
                    safe_runs = [r for r in runs if _is_safe_run(r)]
                    if safe_runs:
                        safe_runs[0].text = translated
                        for r in safe_runs[1:]:
                            r.text = ""
                    else:
                        p.add_run(f" {translated} ")
                else:
                    p.text = translated
            except Exception as pe:
                LOG.error(f"Failed to translate paragraph in-place: {pe}")

        def collect_all_docx_paragraphs(d):
            paragraphs = []
            for p in d.paragraphs:
                paragraphs.append(p)
            
            def collect_from_table(table):
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            paragraphs.append(p)
                        for nested_table in cell.tables:
                            collect_from_table(nested_table)
            
            for table in d.tables:
                collect_from_table(table)
                
            for section in d.sections:
                for h_or_f in [
                    section.header, section.first_page_header, section.even_page_header,
                    section.footer, section.first_page_footer, section.even_page_footer
                ]:
                    if h_or_f:
                        for p in h_or_f.paragraphs:
                            paragraphs.append(p)
                        for table in h_or_f.tables:
                            collect_from_table(table)
            return paragraphs

        all_paragraphs = collect_all_docx_paragraphs(doc)
        paragraphs_to_translate = [p for p in all_paragraphs if p.text.strip()]
        total = len(paragraphs_to_translate)

        # Set initial progress
        self.update_state(
            state='PROGRESS',
            meta={
                'current': 0,
                'total': total,
                'percent': 0
            }
        )

        for idx, p in enumerate(paragraphs_to_translate):
            translate_paragraph_in_place(p, source_lang, target_lang, engine, glossary=glossary)
            current_count = idx + 1
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': current_count,
                    'total': total,
                    'percent': int(100 * current_count / total) if total > 0 else 0
                }
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        try:
            doc.save(str(tmp_path))
            storage.save(trans_key, tmp_path)
            
            # Consume credits after saving successfully
            if creator:
                consume_translation_credits_for_user(creator, estimated_pages)
        finally:
            if tmp_path.exists():
                tmp_path.unlink()

        return {
            'status': 'SUCCESS',
            'current': total,
            'total': total,
            'percent': 100,
            'docx_id': docx_id
        }
 