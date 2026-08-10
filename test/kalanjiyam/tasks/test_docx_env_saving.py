import tempfile
import uuid
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
from docx import Document

from kalanjiyam import database as db
from kalanjiyam import queries as q
from kalanjiyam.utils.storage import MemoryStorage, docx_upload_key
from kalanjiyam.tasks.translation import run_docx_translation


@pytest.fixture
def dummy_docx_file():
    doc = Document()
    doc.add_heading("Test Direct DOCX Document", level=1)
    doc.add_paragraph("This is a paragraph in Sanskrit or English text.")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = Path(tmp.name)
    yield tmp_path
    if tmp_path.exists():
        tmp_path.unlink()


def test_docx_translation_env_saving_disabled(flask_app, dummy_docx_file, monkeypatch):
    monkeypatch.setenv("SAVE_DOCX_DIRECT_TR_DATA", "false")
    docx_id = str(uuid.uuid4())

    storage = MemoryStorage()
    with open(dummy_docx_file, "rb") as f:
        storage.save(docx_upload_key(docx_id), f)

    with patch("kalanjiyam.utils.storage.get_storage", return_value=storage), \
         patch("kalanjiyam.utils.docx_db_saver.save_original_docx_data_to_db") as mock_save, \
         patch("kalanjiyam.utils.translation_engine.translate_text") as mock_translate, \
         patch.object(run_docx_translation, "update_state"), \
         patch("redis.Redis.from_url"):

        mock_resp = MagicMock()
        mock_resp.translated_text = "Translated paragraph text"
        mock_translate.return_value = mock_resp

        result = run_docx_translation.apply(
            kwargs={
                "app_env": "testing",
                "docx_id": docx_id,
                "source_lang": "en",
                "target_lang": "sa",
                "engine": "indictrans2"
            },
            task_id="test-task-id-1"
        ).get()

        assert result["status"] == "SUCCESS"
        mock_save.assert_not_called()


def test_docx_translation_env_saving_enabled(flask_app, dummy_docx_file, monkeypatch):
    monkeypatch.setenv("SAVE_DOCX_DIRECT_TR_DATA", "true")
    docx_id = str(uuid.uuid4())

    storage = MemoryStorage()
    with open(dummy_docx_file, "rb") as f:
        storage.save(docx_upload_key(docx_id), f)

    with patch("kalanjiyam.utils.storage.get_storage", return_value=storage), \
         patch("kalanjiyam.utils.translation_engine.translate_text") as mock_translate, \
         patch.object(run_docx_translation, "update_state"), \
         patch("redis.Redis.from_url"):

        mock_resp = MagicMock()
        mock_resp.translated_text = "Translated paragraph text"
        mock_translate.return_value = mock_resp

        result = run_docx_translation.apply(
            kwargs={
                "app_env": "testing",
                "docx_id": docx_id,
                "source_lang": "en",
                "target_lang": "sa",
                "engine": "indictrans2"
            },
            task_id="test-task-id-2"
        ).get()

        assert result["status"] == "SUCCESS"

        # Verify saved .json.gz payload in storage under docx/saved/{docx_id}.json.gz
        from kalanjiyam.utils.storage import docx_saved_data_key
        saved_key = docx_saved_data_key(docx_id)
        assert storage.exists(saved_key)
        payload = storage.load_json_gz(saved_key)
        assert payload is not None
        assert payload["docx_id"] == docx_id
        assert len(payload["pages"]) >= 1
