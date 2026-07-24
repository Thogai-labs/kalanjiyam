import io
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Inches

from kalanjiyam.tasks.projects import _segment_docx, _extract_docx_images
from kalanjiyam.utils.proofing_utils import documents_to_docx
from kalanjiyam.utils.storage import MemoryStorage


class DummyPage:
    def __init__(self, content, revisions):
        self.revisions = revisions


class DummyRevision:
    def __init__(self, content, document):
        self.content = content
        self.document = document


def test_docx_segmentation_and_export():
    # 1. Create a dummy DOCX document with text formatting, lists, tables, and page breaks
    doc = Document()
    
    # Page 1: Heading, Paragraph with Bold/Italic, Hyperlink, and Bullet List
    doc.add_heading("Sanskrit Text Project", level=1)
    
    p = doc.add_paragraph()
    r1 = p.add_run("This is ")
    r2 = p.add_run("bold")
    r2.bold = True
    r3 = p.add_run(" and ")
    r4 = p.add_run("italic")
    r4.italic = True
    r5 = p.add_run(" text.")

    # Add Bullet List
    doc.add_paragraph("Item 1", style="List Bullet")
    doc.add_paragraph("Item 2", style="List Bullet")

    # Add Math oMath XML directly to paragraph
    p_math = doc.add_paragraph()
    omath_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>E=mc²</m:t></m:r></m:oMath>'
    p_math._p.append(parse_xml(omath_xml.encode('utf-8')))

    # Add page break
    doc.add_page_break()

    # Page 2: Table
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].text = "Header A"
    hdr_cells[1].paragraphs[0].text = "Header B"
    
    row_cells = table.rows[1].cells
    row_cells[0].paragraphs[0].text = "Value A"
    row_cells[1].paragraphs[0].text = "Value B"

    # Save to a temporary file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
        doc.save(tmp_path)

    try:
        # Load doc and test segmentation
        test_doc = Document(tmp_path)
        storage = MemoryStorage()
        image_mapping = _extract_docx_images(test_doc, "test-project", storage)
        
        pages = _segment_docx(test_doc, "test-project", image_mapping)
        
        # We expect exactly 2 pages split at the hard page break
        assert len(pages) == 2
        
        page1_text, page1_html = pages[0]
        page2_text, page2_html = pages[1]
        
        # Verify page 1 content
        assert "Sanskrit Text Project" in page1_text
        assert "bold" in page1_text
        assert "Item 1" in page1_text
        assert "E=mc²" in page1_text
        assert "<h1>Sanskrit Text Project</h1>" in page1_html
        assert "<strong>bold</strong>" in page1_html
        assert "<ul><li>Item 1</li><li>Item 2</li></ul>" in page1_html
        assert 'class="math-placeholder"' in page1_html
        assert 'data-xml=' in page1_html

        # Verify page 2 content
        assert "Header A" in page2_text
        assert "Value B" in page2_text
        assert "<table" in page2_html
        assert "<td>" in page2_html

        # 2. Test compiling back to DOCX
        rev1 = DummyRevision(
            content=page1_text,
            document={
                "content_format": "html",
                "blocks": [{
                    "id": "b1",
                    "type": "paragraph",
                    "content": page1_html + '<p>Please visit <a href="https://example.com">Google</a></p>',
                    "reading_order": 1
                }]
            }
        )
        rev2 = DummyRevision(
            content=page2_text,
            document={
                "content_format": "html",
                "blocks": [{
                    "id": "b2",
                    "type": "paragraph",
                    "content": page2_html + '<div class="docx-column-section" style="column-count: 2;"><p>Col 1 Text</p><p>Col 2 Text</p></div>',
                    "reading_order": 1
                }]
            }
        )
        
        dummy_pages = [
            DummyPage(content=page1_text, revisions=[rev1]),
            DummyPage(content=page2_text, revisions=[rev2])
        ]
        
        compiled_bytes = documents_to_docx(dummy_pages)
        assert len(compiled_bytes) > 0
        
        # Verify compiled file loads as a valid DOCX and has correct elements
        compiled_stream = io.BytesIO(compiled_bytes)
        compiled_doc = Document(compiled_stream)
        
        # Check text exists in paragraphs
        text_content = [p.text for p in compiled_doc.paragraphs]
        assert any("Sanskrit Text Project" in text for text in text_content)
        assert any("bold" in text for text in text_content)
        assert any("Item 1" in text for text in text_content)
        assert any("Google" in text for text in text_content)
        assert any("Col 1 Text" in text for text in text_content)
        assert any("Col 2 Text" in text for text in text_content)
        
        # Check that hyperlink relationships are present
        rels = compiled_doc.part.rels
        hyperlink_urls = [rel.target_ref for rel in rels.values() if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"]
        assert "https://example.com" in hyperlink_urls
        
        # Check that column layouts are preserved
        sections = compiled_doc.sections
        assert len(sections) > 2
        from docx.oxml.ns import qn
        has_2_cols = False
        for sec in sections:
            cols_el = sec._sectPr.find(qn('w:cols'))
            if cols_el is not None and cols_el.get(qn('w:num')) == '2':
                has_2_cols = True
                break
        assert has_2_cols
        
        # Check table cells
        table_cells_text = [cell.text for t in compiled_doc.tables for r in t.rows for cell in r.cells]
        assert "Header A" in table_cells_text
        assert "Value B" in table_cells_text

    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def test_academic_structural_elements():
    # Test footnotes, fonts, and indents roundtrip compiling
    html_content = (
        '<p style="margin-left: 1.5000in; margin-right: 0.5000in; margin-top: 10.00pt; margin-bottom: 12.00pt; line-height: 1.5;">'
        'This is a structured paragraph with '
        '<span style="font-family: Sanskrit2003; font-size: 14.50pt;">Sanskrit script</span> '
        'and a footnote reference '
        '<span class="docx-footnote" data-footnote-id="999" data-footnote-text="This is the footnote details.">1</span>.'
        '</p>'
    )
    
    rev = DummyRevision(
        content="This is plain content.",
        document={
            "content_format": "html",
            "blocks": [{
                "id": "b_acad",
                "type": "paragraph",
                "content": html_content,
                "reading_order": 1
            }]
        }
    )
    
    pages = [DummyPage(content="This is plain content.", revisions=[rev])]
    compiled_bytes = documents_to_docx(pages)
    assert len(compiled_bytes) > 0
    
    # Load document and check elements
    doc = Document(io.BytesIO(compiled_bytes))
    
    # Verify paragraph format
    p = doc.paragraphs[0]
    assert p.paragraph_format.left_indent.inches == 1.5
    assert p.paragraph_format.right_indent.inches == 0.5
    assert p.paragraph_format.space_before.pt == 10.0
    assert p.paragraph_format.space_after.pt == 12.0
    assert p.paragraph_format.line_spacing == 1.5
    
    # Verify run font family and size
    runs = p.runs
    sanskrit_run = None
    for r in runs:
        if r.text == "Sanskrit script":
            sanskrit_run = r
            break
    assert sanskrit_run is not None
    assert sanskrit_run.font.name == "Sanskrit2003"
    assert sanskrit_run.font.size.pt == 14.5
    
    # Verify footnote exists in the document package
    footnotes_part = None
    from docx.opc.packuri import PackURI
    for part in doc.part.package.parts:
        if part.partname == PackURI('/word/footnotes.xml'):
            footnotes_part = part
            break
            
    assert footnotes_part is not None
    from docx.oxml import parse_xml
    footnotes_el = parse_xml(footnotes_part.blob)
    footnote_999 = footnotes_el.find('.//w:footnote[@w:id="999"]', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    assert footnote_999 is not None
    
    footnote_text = "".join(t.text for t in footnote_999.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) if t.text)
    assert "This is the footnote details." in footnote_text

