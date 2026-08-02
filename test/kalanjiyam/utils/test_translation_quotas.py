import pytest
from werkzeug.exceptions import HTTPException

import kalanjiyam.database as db
from kalanjiyam.queries import get_session
from kalanjiyam.utils.quotas import (
    ensure_translation_quota_for_project,
    consume_translation_credit_for_project,
)

def test_translation_quota_enforcement(client):
    session = get_session()
    
    # 1. Setup test organization, users and project
    org = db.Group(
        name="Test Org Quotas",
        slug="test-org-quotas",
        translation_credit_limit=10,
        default_user_translation_limit=5,
    )
    session.add(org)
    session.flush()

    user = db.User(
        username="quota-user",
        email="quota-user@siddhasagaram.in",
        organization_id=org.id,
    )
    user.set_password("quota-password")
    session.add(user)
    session.flush()

    board = db.Board(title="Quota Board")
    session.add(board)
    session.flush()

    project = db.Project(
        slug="quota-project",
        display_title="Quota Project",
        creator_id=user.id,
        board_id=board.id,
    )
    # Associate project with organization
    project.groups.append(org)
    session.add(project)
    session.flush()
    session.commit()

    try:
        # Check quota when within limits (both user and org limits are 0 initially)
        ensure_translation_quota_for_project(project)  # Should not raise any error

        # 2. Test user limit enforcement
        user.translation_credits_used = 5
        session.add(user)
        session.commit()

        # Should raise 402 Payment Required for personal limit
        with pytest.raises(HTTPException) as exc_info:
            ensure_translation_quota_for_project(project)
        assert exc_info.value.code == 402
        assert "personal Translation credit limit" in exc_info.value.description

        # Reset user credits, test org limit enforcement
        user.translation_credits_used = 0
        org.translation_credits_used = 10
        session.add(user)
        session.add(org)
        session.commit()

        # Should raise 402 Payment Required for organization limit
        with pytest.raises(HTTPException) as exc_info:
            ensure_translation_quota_for_project(project)
        assert exc_info.value.code == 402
        assert "Organization/Tenant Translation credits" in exc_info.value.description

        # 3. Test consume credits
        user.translation_credits_used = 2
        org.translation_credits_used = 4
        session.add(user)
        session.add(org)
        session.commit()

        consume_translation_credit_for_project(project)

        session.refresh(user)
        session.refresh(org)
        assert user.translation_credits_used == 3
        assert org.translation_credits_used == 5

    finally:
        # Cleanup
        session.delete(project)
        session.delete(board)
        session.delete(user)
        session.delete(org)
        session.commit()


def test_docx_translation_quota_and_estimation(client):
    from docx import Document
    from kalanjiyam.utils.quotas import (
        estimate_docx_pages,
        ensure_translation_quota_for_user,
        consume_translation_credits_for_user,
    )
    
    # 1. Test estimate_docx_pages
    doc = Document()
    doc.add_paragraph("Paragraph 1 of content")
    doc.add_page_break()
    doc.add_paragraph("Paragraph 2 of content")
    # Add a paragraph with >1500 characters to trigger second page threshold
    long_text = "A" * 1501
    doc.add_paragraph(long_text)
    # Add paragraph 3 which should go on page 3 since page 2 was flushed
    doc.add_paragraph("Paragraph 3 of content")
    
    pages = estimate_docx_pages(doc)
    # page 1: Paragraph 1 (ends with page break -> page count + 1)
    # page 2: Paragraph 2 + Paragraph 3 (Paragraph 3 has > 1500 chars -> page count + 1)
    # page 3: Paragraph 4 -> (remaining page -> page count + 1)
    # total should be 3 pages
    assert pages == 3

    # 2. Test ensure_translation_quota_for_user and consume_translation_credits_for_user
    session = get_session()
    org = db.Group(
        name="Docx Org Quotas",
        slug="docx-org-quotas",
        translation_credit_limit=10,
        default_user_translation_limit=5,
    )
    session.add(org)
    session.flush()

    user = db.User(
        username="docx-quota-user",
        email="docx-quota-user@siddhasagaram.in",
        organization_id=org.id,
    )
    user.set_password("quota-password")
    session.add(user)
    session.flush()
    session.commit()

    try:
        # Should not raise any error for 2 pages (requires 2 credits, 0 used)
        ensure_translation_quota_for_user(user, required_credits=2)

        # Exceed personal limit
        with pytest.raises(HTTPException) as exc_info:
            ensure_translation_quota_for_user(user, required_credits=6)
        assert exc_info.value.code == 402
        assert "personal Translation credit limit" in exc_info.value.description

        # Exceed org limit
        org.translation_credits_used = 9
        session.add(org)
        session.commit()
        with pytest.raises(HTTPException) as exc_info:
            ensure_translation_quota_for_user(user, required_credits=2)
        assert exc_info.value.code == 402
        assert "Organization/Tenant Translation credits" in exc_info.value.description

        # Reset limits and consume
        org.translation_credits_used = 1
        user.translation_credits_used = 1
        session.add(org)
        session.add(user)
        session.commit()

        consume_translation_credits_for_user(user, credits=3)
        session.refresh(user)
        session.refresh(org)
        assert user.translation_credits_used == 4
        assert org.translation_credits_used == 4

    finally:
        session.delete(user)
        session.delete(org)
        session.commit()

